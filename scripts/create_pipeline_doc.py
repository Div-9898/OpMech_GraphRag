#!/usr/bin/env python3
"""Generate comprehensive pipeline documentation as Word document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def create_pipeline_document():
    """Create the MoE Graph Builder Pipeline Documentation."""
    doc = Document()

    # Set document title
    title = doc.add_heading('MoE Graph Builder: Complete Pipeline Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph('Knowledge Graph Construction from Apple SEC Filings')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture Overview',
        '3. Phase 1: Data Ingestion Pipeline',
        '   3.1 SEC EDGAR Fetcher',
        '   3.2 HTML Parser',
        '   3.3 XBRL Processor',
        '   3.4 Embedding Engine',
        '4. Phase 2: Expert Modules (Mixture-of-Experts)',
        '   4.1 CrossReferenceHunter',
        '   4.2 CausalChainBuilder',
        '   4.3 TemporalLinker',
        '   4.4 TableTextConnector',
        '   4.5 SemanticBridge',
        '   4.6 EntityExtractor',
        '5. Phase 3: Graph Construction',
        '   5.1 Graph Builder Orchestration',
        '   5.2 Connectivity Enforcer',
        '   5.3 Neo4j Export',
        '6. Phase 4: Evaluation Framework',
        '7. Data Flow Diagram',
        '8. Technical Specifications',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number' if item[0].isdigit() and '.' not in item[1:3] else 'List Bullet')

    doc.add_page_break()

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'The MoE (Mixture-of-Experts) Graph Builder is a sophisticated knowledge graph construction system '
        'designed to extract, analyze, and interconnect information from Apple\'s SEC filings. The system '
        'processes 12 quarterly and annual reports (10-K and 10-Q filings) from fiscal years 2022-2024, '
        'creating a rich, interconnected knowledge graph with approximately 1,200+ nodes and 15,000+ edges.'
    )

    doc.add_paragraph(
        'Key Capabilities:',
        style='Heading 3'
    )
    capabilities = [
        'Automated SEC filing retrieval with rate limiting and caching',
        'Multi-modal document parsing (HTML, XBRL, tables)',
        'Financial domain-specific embeddings using FinBERT',
        'Six specialized expert modules for relationship extraction',
        'LLM-powered analysis using Qwen2.5-7B via vLLM',
        'Neo4j graph database for storage and querying',
        'Guaranteed single connected component through bridge edges',
    ]
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')

    # 2. System Architecture
    doc.add_heading('2. System Architecture Overview', level=1)

    doc.add_paragraph(
        'The system follows a modular, pipeline-based architecture with four main phases:'
    )

    arch_table = doc.add_table(rows=5, cols=3)
    arch_table.style = 'Table Grid'
    headers = ['Phase', 'Components', 'Output']
    for i, header in enumerate(headers):
        arch_table.rows[0].cells[i].text = header
        arch_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    arch_data = [
        ('1. Data Ingestion', 'SEC Fetcher, HTML Parser, XBRL Processor, Embedding Engine', 'Nodes with embeddings'),
        ('2. Expert Processing', '6 MoE Experts (Cross-Ref, Causal, Temporal, Table-Text, Semantic, Entity)', 'Candidate edges'),
        ('3. Graph Construction', 'Graph Builder, Connectivity Enforcer, Neo4j Client', 'Connected knowledge graph'),
        ('4. Evaluation', 'Gold Standard Generator, Metrics Calculator', 'Performance metrics'),
    ]
    for i, (phase, components, output) in enumerate(arch_data, 1):
        arch_table.rows[i].cells[0].text = phase
        arch_table.rows[i].cells[1].text = components
        arch_table.rows[i].cells[2].text = output

    doc.add_paragraph()

    # Technology Stack
    doc.add_heading('Technology Stack', level=2)
    tech_items = [
        ('Language', 'Python 3.11+'),
        ('ML Framework', 'PyTorch, Transformers, Sentence-Transformers'),
        ('Embeddings', 'FinBERT (ProsusAI/finbert) - 768 dimensions'),
        ('LLM', 'Qwen2.5-7B-Instruct via vLLM (OpenAI-compatible API)'),
        ('Graph Database', 'Neo4j 5.x'),
        ('Data Validation', 'Pydantic v2'),
        ('Logging', 'Loguru'),
    ]
    for tech, value in tech_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{tech}: ').bold = True
        p.add_run(value)

    doc.add_page_break()

    # 3. Phase 1: Data Ingestion
    doc.add_heading('3. Phase 1: Data Ingestion Pipeline', level=1)

    doc.add_paragraph(
        'The data ingestion pipeline is responsible for fetching SEC filings, parsing their content, '
        'and generating embeddings for semantic analysis. This phase transforms raw regulatory documents '
        'into structured nodes ready for relationship extraction.'
    )

    # 3.1 SEC Fetcher
    doc.add_heading('3.1 SEC EDGAR Fetcher', level=2)
    doc.add_paragraph('File: src/ingestion/sec_fetcher.py', style='Intense Quote')

    doc.add_paragraph(
        'The SEC Fetcher retrieves Apple\'s 10-K (annual) and 10-Q (quarterly) filings from the SEC EDGAR database. '
        'It implements robust error handling, rate limiting, and caching mechanisms.'
    )

    doc.add_heading('Key Features:', level=3)
    features = [
        'Rate Limiting: Respects SEC\'s 10 requests/second limit with automatic throttling',
        'Exponential Backoff: Retries failed requests with increasing delays (1s, 2s, 4s...)',
        'Caching: Stores raw filings in data/raw/{CIK}/{accession_number}/',
        'Filing Discovery: Uses SEC\'s submissions API to find all available filings',
        'Document Selection: Identifies primary filing document from filing index',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('Target Filings:', level=3)
    filings_table = doc.add_table(rows=4, cols=3)
    filings_table.style = 'Table Grid'
    filings_table.rows[0].cells[0].text = 'Filing Type'
    filings_table.rows[0].cells[1].text = 'Periods'
    filings_table.rows[0].cells[2].text = 'Count'
    for cell in filings_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    filings_table.rows[1].cells[0].text = '10-K (Annual)'
    filings_table.rows[1].cells[1].text = 'FY2022, FY2023, FY2024'
    filings_table.rows[1].cells[2].text = '3'
    filings_table.rows[2].cells[0].text = '10-Q (Quarterly)'
    filings_table.rows[2].cells[1].text = 'Q1-Q3 for 2022, 2023, 2024'
    filings_table.rows[2].cells[2].text = '9'
    filings_table.rows[3].cells[0].text = 'Total'
    filings_table.rows[3].cells[1].text = ''
    filings_table.rows[3].cells[2].text = '12 filings'

    doc.add_paragraph()

    # 3.2 HTML Parser
    doc.add_heading('3.2 HTML Parser', level=2)
    doc.add_paragraph('File: src/ingestion/html_parser.py', style='Intense Quote')

    doc.add_paragraph(
        'The HTML Parser extracts structured content from SEC filing HTML documents, identifying '
        'different sections, tables, and notes while preserving document hierarchy.'
    )

    doc.add_heading('Extraction Components:', level=3)
    components = [
        ('Text Sections', 'Item 1 (Business), Item 1A (Risk Factors), Item 7 (MD&A), Item 7A (Market Risk)'),
        ('Notes', 'Footnotes to financial statements (Note 1 through Note N)'),
        ('Tables', 'Financial tables with numeric data'),
    ]
    for comp, desc in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{comp}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Chunking Strategy:', level=3)
    doc.add_paragraph(
        'Large text sections are split into manageable chunks for embedding generation:'
    )
    chunk_items = [
        'Default chunk size: 5,000 characters',
        'MD&A sections (Item 7): 8,000 characters (requires more context)',
        'Risk Factors (Item 1A): 8,000 characters',
        'Minimum paragraph length: 50 characters',
        'Chunks are split at paragraph boundaries when possible',
    ]
    for item in chunk_items:
        doc.add_paragraph(item, style='List Bullet')

    # 3.3 XBRL Processor
    doc.add_heading('3.3 XBRL Processor', level=2)
    doc.add_paragraph('File: src/ingestion/xbrl_processor.py', style='Intense Quote')

    doc.add_paragraph(
        'The XBRL (eXtensible Business Reporting Language) Processor extracts structured financial data '
        'from SEC\'s Company Facts API, providing standardized financial metrics with temporal context.'
    )

    doc.add_heading('Data Extracted:', level=3)
    xbrl_items = [
        ('Financial Tags', 'us-gaap:Revenues, us-gaap:NetIncomeLoss, us-gaap:Assets, etc.'),
        ('Values', 'Numeric values with units (USD, shares, etc.)'),
        ('Periods', 'Fiscal year/quarter with start and end dates'),
        ('Context', 'Segment information, geographic breakdowns'),
    ]
    for item, desc in xbrl_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}: ').bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'Each XBRL fact becomes a FINANCIAL_LINE node with metadata including the XBRL tag, '
        'value, unit, and time period for temporal linking.'
    )

    # 3.4 Embedding Engine
    doc.add_heading('3.4 Embedding Engine', level=2)
    doc.add_paragraph('File: src/ingestion/embedding_engine.py', style='Intense Quote')

    doc.add_paragraph(
        'The Embedding Engine generates dense vector representations of text using FinBERT, '
        'a BERT model fine-tuned on financial text for superior domain-specific understanding.'
    )

    doc.add_heading('Model Details:', level=3)
    model_items = [
        ('Model', 'ProsusAI/finbert'),
        ('Embedding Dimension', '768'),
        ('Max Sequence Length', '512 tokens'),
        ('Device', 'CUDA (GPU-accelerated)'),
        ('Batch Size', '32'),
    ]
    for item, value in model_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}: ').bold = True
        p.add_run(value)

    doc.add_heading('Processing Pipeline:', level=3)
    doc.add_paragraph('1. Text Preprocessing: Truncate/chunk to 512 tokens')
    doc.add_paragraph('2. Tokenization: Convert text to token IDs using FinBERT tokenizer')
    doc.add_paragraph('3. Encoding: Pass through FinBERT model with mean pooling')
    doc.add_paragraph('4. Normalization: L2 normalize embeddings for cosine similarity')
    doc.add_paragraph('5. Caching: Store in NPZ format at data/embeddings/{filing_id}/')

    doc.add_page_break()

    # 4. Phase 2: Expert Modules
    doc.add_heading('4. Phase 2: Expert Modules (Mixture-of-Experts)', level=1)

    doc.add_paragraph(
        'The Mixture-of-Experts (MoE) architecture employs six specialized expert modules, each designed '
        'to identify specific types of relationships within and across SEC filings. Each expert can operate '
        'in two modes: rule-based (fast, deterministic) and LLM-enhanced (more accurate, slower).'
    )

    # Expert Summary Table
    doc.add_heading('Expert Summary:', level=2)
    expert_table = doc.add_table(rows=7, cols=4)
    expert_table.style = 'Table Grid'
    headers = ['Expert', 'Edge Types', 'Primary Method', 'LLM Enhancement']
    for i, h in enumerate(headers):
        expert_table.rows[0].cells[i].text = h
        expert_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    experts = [
        ('CrossReferenceHunter', 'REFERS_TO', 'Regex patterns', 'Reference validation'),
        ('CausalChainBuilder', 'CAUSED_BY, LEADS_TO', 'Dependency parsing', 'Full causal extraction'),
        ('TemporalLinker', 'TEMPORAL_NEXT', 'XBRL tag matching', 'Temporal validation'),
        ('TableTextConnector', 'DISCUSSES, EXPLAINS_LINE_ITEM', 'Numeric matching', 'Context linking'),
        ('SemanticBridge', 'SEMANTICALLY_SIMILAR', 'Embedding similarity', 'Relationship classification'),
        ('EntityExtractor', 'MENTIONS_ENTITY, ENTITY_RELATED_TO', 'N/A (LLM only)', 'Entity extraction'),
    ]
    for i, (name, edges, method, llm) in enumerate(experts, 1):
        expert_table.rows[i].cells[0].text = name
        expert_table.rows[i].cells[1].text = edges
        expert_table.rows[i].cells[2].text = method
        expert_table.rows[i].cells[3].text = llm

    doc.add_paragraph()

    # 4.1 CrossReferenceHunter
    doc.add_heading('4.1 CrossReferenceHunter', level=2)
    doc.add_paragraph('File: src/experts/cross_reference.py', style='Intense Quote')

    doc.add_paragraph(
        'Identifies explicit cross-references within SEC filings, such as references to specific Notes, '
        'Items, or other sections of the document.'
    )

    doc.add_heading('Detection Patterns:', level=3)
    patterns = [
        '"See Note 3" → Links to Note 3 - Revenue Recognition',
        '"refer to Item 7" → Links to MD&A section',
        '"as discussed in Part II, Item 8" → Links to financial statements',
        '"the following table" → Links to adjacent table',
    ]
    for p in patterns:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('Algorithm:', level=3)
    doc.add_paragraph(
        '1. Regex Matching: Scan text for reference patterns (Note X, Item Y, Part Z)\n'
        '2. Target Resolution: Map references to actual document sections\n'
        '3. LLM Validation (optional): Confirm reference validity and extract context\n'
        '4. Confidence Scoring: Higher for exact matches, lower for implicit references'
    )

    # 4.2 CausalChainBuilder
    doc.add_heading('4.2 CausalChainBuilder', level=2)
    doc.add_paragraph('File: src/experts/causal.py', style='Intense Quote')

    doc.add_paragraph(
        'Extracts cause-and-effect relationships from financial narrative text, identifying '
        'factors that drive business outcomes.'
    )

    doc.add_heading('Causal Indicators:', level=3)
    indicators = [
        'Forward causation: "due to", "because of", "as a result of", "driven by"',
        'Backward causation: "led to", "resulted in", "caused", "contributed to"',
        'Conditional: "if...then", "when...consequently"',
    ]
    for ind in indicators:
        doc.add_paragraph(ind, style='List Bullet')

    doc.add_heading('LLM-Enhanced Extraction:', level=3)
    doc.add_paragraph(
        'When LLM mode is enabled, the expert uses Qwen2.5-7B to:\n'
        '- Extract complex causal chains spanning multiple sentences\n'
        '- Identify implicit causal relationships\n'
        '- Classify causal direction (forward/backward)\n'
        '- Assign confidence scores based on linguistic evidence'
    )

    # 4.3 TemporalLinker
    doc.add_heading('4.3 TemporalLinker', level=2)
    doc.add_paragraph('File: src/experts/temporal.py', style='Intense Quote')

    doc.add_paragraph(
        'Creates temporal connections between corresponding sections across different time periods, '
        'enabling year-over-year and quarter-over-quarter analysis.'
    )

    doc.add_heading('Linking Strategies:', level=3)
    strategies = [
        ('XBRL Tag Matching', 'Same financial metric (e.g., us-gaap:Revenues) across periods'),
        ('Note Number Matching', 'Same note number across filings (Note 3 FY2023 → Note 3 FY2024)'),
        ('Section Matching', 'Same Item/Part across filings (Item 7 FY2023 → Item 7 FY2024)'),
        ('Embedding Similarity', 'High cosine similarity (>0.90) between text chunks'),
    ]
    for strategy, desc in strategies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{strategy}: ').bold = True
        p.add_run(desc)

    # 4.4 TableTextConnector
    doc.add_heading('4.4 TableTextConnector', level=2)
    doc.add_paragraph('File: src/experts/table_text.py', style='Intense Quote')

    doc.add_paragraph(
        'Links financial data (tables, XBRL facts) to explanatory narrative text, '
        'connecting quantitative and qualitative information.'
    )

    doc.add_heading('Connection Methods:', level=3)
    methods = [
        ('Numeric Matching', 'Find text mentioning same values (5% tolerance for rounding)'),
        ('Embedding Similarity', 'Match FINANCIAL_LINE nodes to similar TEXT_SECTION nodes (>0.80)'),
        ('XBRL-Note Linking', 'Connect XBRL tags to corresponding Note sections'),
        ('LLM Context Linking', 'Use LLM to identify explanatory relationships'),
    ]
    for method, desc in methods:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{method}: ').bold = True
        p.add_run(desc)

    # 4.5 SemanticBridge
    doc.add_heading('4.5 SemanticBridge', level=2)
    doc.add_paragraph('File: src/experts/semantic.py', style='Intense Quote')

    doc.add_paragraph(
        'Creates semantic similarity edges between conceptually related content, '
        'both within and across filings.'
    )

    doc.add_heading('Similarity Thresholds:', level=3)
    thresholds = [
        ('Within-filing similarity', '> 0.85 cosine similarity'),
        ('Cross-filing similarity', '> 0.85 cosine similarity'),
        ('Bridge edges (connectivity)', '> 0.70 cosine similarity'),
    ]
    for thresh, value in thresholds:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{thresh}: ').bold = True
        p.add_run(value)

    # 4.6 EntityExtractor
    doc.add_heading('4.6 EntityExtractor', level=2)
    doc.add_paragraph('File: src/experts/entity_extractor.py', style='Intense Quote')

    doc.add_paragraph(
        'Extracts named entities from SEC filings using LLM, creating entity nodes and '
        'linking them to source documents. This expert requires LLM mode to function.'
    )

    doc.add_heading('Entity Types:', level=3)
    entities = [
        ('COMPANY', 'Apple, subsidiaries, suppliers, competitors'),
        ('PRODUCT', 'iPhone, Mac, iPad, Services, AppleCare'),
        ('SEGMENT', 'Americas, Europe, Greater China, Japan'),
        ('PERSON', 'Tim Cook, executives mentioned by name'),
        ('FINANCIAL_METRIC', 'Revenue, Net Income, Gross Margin'),
        ('RISK_FACTOR', 'Identified business risks'),
        ('REGULATION', 'GAAP, ASC 606, GDPR'),
    ]
    for etype, examples in entities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{etype}: ').bold = True
        p.add_run(examples)

    doc.add_page_break()

    # 5. Phase 3: Graph Construction
    doc.add_heading('5. Phase 3: Graph Construction', level=1)

    # 5.1 Graph Builder
    doc.add_heading('5.1 Graph Builder Orchestration', level=2)
    doc.add_paragraph('File: src/graph/builder.py', style='Intense Quote')

    doc.add_paragraph(
        'The Graph Builder orchestrates the entire graph construction process, coordinating '
        'expert execution, edge merging, and connectivity enforcement.'
    )

    doc.add_heading('Build Process:', level=3)
    steps = [
        '1. Load all nodes and embeddings from data ingestion phase',
        '2. Run Entity Extraction (if enabled) to create entity nodes',
        '3. Execute each expert module sequentially to discover edges',
        '4. Merge and deduplicate edges (keep highest confidence)',
        '5. Verify connectivity and add bridge edges if needed',
        '6. Export to Neo4j and save to files',
    ]
    for step in steps:
        doc.add_paragraph(step)

    doc.add_heading('Edge Deduplication:', level=3)
    doc.add_paragraph(
        'When multiple experts discover the same edge, the system keeps only the edge with '
        'the highest confidence score. Edge identity is determined by (source_id, target_id, edge_type) tuple.'
    )

    # 5.2 Connectivity Enforcer
    doc.add_heading('5.2 Connectivity Enforcer', level=2)
    doc.add_paragraph('File: src/graph/connectivity.py', style='Intense Quote')

    doc.add_paragraph(
        'Ensures the final graph is a single connected component by adding bridge edges '
        'between disconnected subgraphs.'
    )

    doc.add_heading('Algorithm:', level=3)
    algo_steps = [
        '1. Detect connected components using BFS traversal',
        '2. If multiple components exist, identify pairs of nodes to connect',
        '3. For each component pair, find the most similar nodes (by embedding)',
        '4. Create BRIDGE edges with similarity-based confidence',
        '5. Repeat until single component achieved',
    ]
    for step in algo_steps:
        doc.add_paragraph(step)

    doc.add_heading('Bridge Edge Properties:', level=3)
    doc.add_paragraph(
        '- Edge type: SEMANTICALLY_SIMILAR\n'
        '- Expert: ConnectivityEnforcer\n'
        '- Confidence: Based on embedding cosine similarity (minimum 0.70)\n'
        '- Evidence: "Bridge edge for connectivity"'
    )

    # 5.3 Neo4j Export
    doc.add_heading('5.3 Neo4j Export', level=2)
    doc.add_paragraph('File: src/graph/neo4j_client.py', style='Intense Quote')

    doc.add_paragraph(
        'The Neo4j Client handles all interactions with the Neo4j graph database, '
        'including batch insertions and index creation.'
    )

    doc.add_heading('Database Schema:', level=3)

    doc.add_paragraph('Node Labels:', style='Heading 4')
    node_labels = [
        'Node (all nodes) with properties: id, type, text, filing_id, period, section, note_number, xbrl_tag, value',
    ]
    for label in node_labels:
        doc.add_paragraph(label, style='List Bullet')

    doc.add_paragraph('Relationship Types:', style='Heading 4')
    rel_types = [
        'REFERS_TO - Cross-references between sections',
        'CAUSED_BY / LEADS_TO - Causal relationships',
        'TEMPORAL_NEXT - Temporal connections across periods',
        'DISCUSSES / EXPLAINS_LINE_ITEM - Table-text connections',
        'SEMANTICALLY_SIMILAR - Semantic similarity edges',
        'MENTIONS_ENTITY / ENTITY_RELATED_TO - Entity relationships',
    ]
    for rel in rel_types:
        doc.add_paragraph(rel, style='List Bullet')

    doc.add_heading('Indexes:', level=3)
    indexes = [
        'Node.id - Unique constraint for node lookup',
        'Node.type - Index for filtering by node type',
        'Node.filing_id - Index for filtering by filing',
    ]
    for idx in indexes:
        doc.add_paragraph(idx, style='List Bullet')

    doc.add_page_break()

    # 6. Phase 4: Evaluation
    doc.add_heading('6. Phase 4: Evaluation Framework', level=1)

    doc.add_paragraph(
        'The evaluation framework measures expert performance against gold standard annotations, '
        'calculating precision, recall, and F1 scores for each expert.'
    )

    doc.add_heading('Gold Standard Generation:', level=2)
    doc.add_paragraph('File: scripts/generate_gold_standard.py', style='Intense Quote')

    doc.add_paragraph(
        'Gold standard datasets are generated using a combination of:\n'
        '- Pattern-based positive samples (explicit references, XBRL matches)\n'
        '- LLM-assisted annotation for ambiguous cases\n'
        '- Heuristic-based negative samples (cross-filing unrelated pairs)'
    )

    doc.add_heading('Metrics Calculated:', level=2)
    metrics = [
        ('Precision', 'TP / (TP + FP) - Accuracy of positive predictions'),
        ('Recall', 'TP / (TP + FN) - Coverage of true positives'),
        ('F1 Score', '2 * (Precision * Recall) / (Precision + Recall)'),
        ('Calibration Error', 'Expected Calibration Error (ECE) for confidence scores'),
    ]
    for metric, desc in metrics:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{metric}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Target F1 Thresholds:', level=2)
    threshold_table = doc.add_table(rows=6, cols=3)
    threshold_table.style = 'Table Grid'
    headers = ['Expert', 'Minimum F1', 'Target F1']
    for i, h in enumerate(headers):
        threshold_table.rows[0].cells[i].text = h
        threshold_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    thresholds = [
        ('CrossReferenceHunter', '0.82', '0.88'),
        ('CausalChainBuilder', '0.72', '0.78'),
        ('TemporalLinker', '0.87', '0.92'),
        ('TableTextConnector', '0.80', '0.85'),
        ('SemanticBridge', '0.70', '0.75'),
    ]
    for i, (expert, min_f1, target) in enumerate(thresholds, 1):
        threshold_table.rows[i].cells[0].text = expert
        threshold_table.rows[i].cells[1].text = min_f1
        threshold_table.rows[i].cells[2].text = target

    doc.add_page_break()

    # 7. Data Flow Diagram
    doc.add_heading('7. Data Flow Diagram', level=1)

    doc.add_paragraph(
        'The following diagram illustrates the complete data flow through the system:'
    )

    flow_text = """
    ┌─────────────────────────────────────────────────────────────────┐
    │                      SEC EDGAR API                               │
    │                  (Apple 10-K, 10-Q filings)                      │
    └─────────────────────────────────────────────────────────────────┘
                                    │
                                    ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                    DATA INGESTION PHASE                          │
    │  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐              │
    │  │ SEC Fetcher │→│ HTML Parser │→│ XBRL Proc.  │              │
    │  └─────────────┘  └─────────────┘  └─────────────┘              │
    │                           │                                      │
    │                           ▼                                      │
    │                  ┌─────────────────┐                             │
    │                  │ Embedding Engine│ (FinBERT)                   │
    │                  └─────────────────┘                             │
    └─────────────────────────────────────────────────────────────────┘
                                    │
                    Nodes + Embeddings (~1,200 nodes)
                                    │
                                    ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                   EXPERT PROCESSING PHASE                        │
    │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐             │
    │  │ CrossRef     │ │ Causal       │ │ Temporal     │             │
    │  │ Hunter       │ │ ChainBuilder │ │ Linker       │             │
    │  └──────────────┘ └──────────────┘ └──────────────┘             │
    │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐             │
    │  │ TableText    │ │ Semantic     │ │ Entity       │             │
    │  │ Connector    │ │ Bridge       │ │ Extractor    │             │
    │  └──────────────┘ └──────────────┘ └──────────────┘             │
    │                           │                                      │
    │                    vLLM (Qwen2.5-7B) ◄────────────────           │
    └─────────────────────────────────────────────────────────────────┘
                                    │
                    Candidate Edges (~16,000 edges)
                                    │
                                    ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                  GRAPH CONSTRUCTION PHASE                        │
    │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐             │
    │  │ Edge         │→│ Connectivity │→│ Neo4j        │             │
    │  │ Deduplication│ │ Enforcer     │ │ Export       │             │
    │  └──────────────┘ └──────────────┘ └──────────────┘             │
    └─────────────────────────────────────────────────────────────────┘
                                    │
                                    ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │                       NEO4J DATABASE                             │
    │              Connected Knowledge Graph                           │
    │           (~1,200 nodes, ~16,000 edges)                          │
    └─────────────────────────────────────────────────────────────────┘
    """

    # Add as preformatted text
    p = doc.add_paragraph()
    run = p.add_run(flow_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    doc.add_page_break()

    # 8. Technical Specifications
    doc.add_heading('8. Technical Specifications', level=1)

    doc.add_heading('Hardware Requirements:', level=2)
    hw_specs = [
        ('GPU', 'NVIDIA RTX 5090 (24GB VRAM) or equivalent'),
        ('RAM', '32GB minimum'),
        ('Storage', '10GB for data and models'),
        ('CPU', 'Modern multi-core processor'),
    ]
    for spec, value in hw_specs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{spec}: ').bold = True
        p.add_run(value)

    doc.add_heading('Software Dependencies:', level=2)
    sw_deps = [
        'Python 3.11+',
        'PyTorch 2.0+ with CUDA support',
        'Transformers 4.36+',
        'vLLM 0.4+ for LLM serving',
        'Neo4j 5.x',
        'Docker (for Neo4j container)',
    ]
    for dep in sw_deps:
        doc.add_paragraph(dep, style='List Bullet')

    doc.add_heading('Performance Metrics:', level=2)
    perf_table = doc.add_table(rows=6, cols=2)
    perf_table.style = 'Table Grid'
    perf_table.rows[0].cells[0].text = 'Metric'
    perf_table.rows[0].cells[1].text = 'Value'
    for cell in perf_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    perf_data = [
        ('Total Nodes', '~1,200'),
        ('Total Edges', '~16,000'),
        ('Connected Components', '1 (guaranteed)'),
        ('Average Node Degree', '~25'),
        ('Build Time (with LLM)', '~45 minutes'),
    ]
    for i, (metric, value) in enumerate(perf_data, 1):
        perf_table.rows[i].cells[0].text = metric
        perf_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    doc.add_heading('Configuration Options:', level=2)
    doc.add_paragraph(
        'Key configuration options are controlled via environment variables (.env file) '
        'and command-line arguments:'
    )

    config_items = [
        ('--use-llm', 'Enable LLM mode for all experts'),
        ('--extract-entities', 'Enable entity extraction using LLM'),
        ('--skip-neo4j', 'Skip Neo4j export (file output only)'),
        ('--experts', 'Comma-separated list of experts to run'),
        ('VLLM_API_BASE', 'vLLM server URL (default: http://localhost:8000/v1)'),
        ('VLLM_MODEL', 'LLM model name (default: Qwen/Qwen2.5-7B-Instruct)'),
    ]
    for opt, desc in config_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{opt}: ').bold = True
        p.add_run(desc)

    # Save document
    output_path = Path('/home/divyansh/AIF_FInal_Project/docs/MoE_Graph_Builder_Pipeline_Documentation.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_pipeline_document()
