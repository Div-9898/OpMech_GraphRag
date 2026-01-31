#!/usr/bin/env python3
"""
Generate comprehensive MOE Architecture Technical Documentation as a Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_document():
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('OpMech-GraphRAG: Mixture of Experts Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Complete Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    # Document info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f'Version 1.0 | Generated: {datetime.now().strftime("%Y-%m-%d")}').italic = True

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1. Executive Summary', 3),
        ('2. System Architecture Overview', 4),
        ('3. Mixture of Experts (MoE) Framework', 6),
        ('   3.1 CrossReferenceHunter Expert', 7),
        ('   3.2 CausalChainBuilder Expert', 9),
        ('   3.3 TemporalLinker Expert', 11),
        ('   3.4 TableTextConnector Expert', 13),
        ('   3.5 SemanticBridge Expert', 15),
        ('   3.6 EntityExtractor Expert', 17),
        ('4. ETL Pipeline for Document Ingestion', 19),
        ('   4.1 SEC Fetcher Component', 20),
        ('   4.2 HTML Parser Component', 21),
        ('   4.3 XBRL Processor Component', 22),
        ('   4.4 Embedding Engine (FinBERT)', 23),
        ('5. Commutator Configuration', 25),
        ('   5.1 Divergence Formula', 25),
        ('   5.2 Divergence Components', 26),
        ('   5.3 Operator Scoring', 28),
        ('6. Operator Configurations', 29),
        ('   6.1 OperatorA (Structure-First)', 29),
        ('   6.2 OperatorB (Narrative-First)', 32),
        ('7. Explore/Exploit Controller', 35),
        ('8. Mode Selection & Trust Decisions', 38),
        ('9. System Integration & Query Flow', 42),
        ('10. Data Classes & Models', 46),
        ('11. Configuration Reference', 50),
        ('12. Appendices', 52),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f'{item}').bold = '.' not in item[:5]
        p.add_run(f' {"." * (60 - len(item))} {page}')

    doc.add_page_break()

    # ========================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ========================================================================
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        'This document provides comprehensive technical documentation for the OpMech-GraphRAG '
        '(Operator Mechanics Graph Retrieval-Augmented Generation) system, an advanced Mixture '
        'of Experts (MoE) architecture designed for financial document analysis. The system '
        'is specifically tailored for processing Apple Inc. SEC filings (10-K and 10-Q reports) '
        'and answering complex financial queries with high accuracy and explainability.'
    )

    doc.add_heading('Key Architectural Components', level=2)

    components = [
        ('Six Specialized Experts',
         'Domain-specific modules that identify different types of relationships in financial documents'),
        ('Dual Operator System',
         'Two parallel operators (Structure-First and Narrative-First) that traverse the knowledge graph from different perspectives'),
        ('Commutator',
         'A mathematical framework measuring divergence between operator outputs to guide explore/exploit decisions'),
        ('ETL Pipeline',
         'Complete data ingestion system for fetching, parsing, and embedding SEC filings'),
        ('Neo4j Knowledge Graph',
         'Graph database storing nodes (document chunks) and edges (relationships) with confidence scores'),
        ('FinBERT Embeddings',
         'Financial domain-specific embeddings for semantic similarity computation'),
    ]

    for title_text, desc in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title_text}: ').bold = True
        p.add_run(desc)

    doc.add_heading('System Objectives', level=2)

    objectives = [
        'Provide accurate answers to complex financial queries using SEC filing data',
        'Maintain full traceability from answers back to source evidence',
        'Handle both quantitative (XBRL-based) and qualitative (narrative) questions',
        'Dynamically balance exploration vs. exploitation based on query complexity',
        'Ensure consistency between different analytical perspectives',
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_page_break()

    # ========================================================================
    # SECTION 2: SYSTEM ARCHITECTURE OVERVIEW
    # ========================================================================
    doc.add_heading('2. System Architecture Overview', level=1)

    doc.add_paragraph(
        'The OpMech-GraphRAG system follows a modular architecture with clear separation of concerns. '
        'The architecture can be divided into three main layers: Data Ingestion, Knowledge Graph, '
        'and Query Processing.'
    )

    doc.add_heading('2.1 High-Level Architecture', level=2)

    # Architecture diagram as text
    arch_text = '''
┌─────────────────────────────────────────────────────────────────────────────┐
│                          OpMech-GraphRAG System                              │
├─────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  ┌─────────────────────────────────────────────────────────────────────────┐│
│  │                      QUERY PROCESSING LAYER                              ││
│  │  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐    ││
│  │  │ Query       │  │OperatorA   │  │OperatorB   │  │ Mode        │    ││
│  │  │ Classifier  │  │(Structure) │  │(Narrative) │  │ Selector    │    ││
│  │  └──────┬──────┘  └──────┬──────┘  └──────┬──────┘  └──────┬──────┘    ││
│  │         │                │                │                │           ││
│  │  ┌──────┴────────────────┴────────────────┴────────────────┴──────┐    ││
│  │  │                      COMMUTATOR                                 │    ││
│  │  │     Δ(q,h) = w_E·Δ_E + w_V·Δ_V + w_A·Δ_A + w_C·Δ_C            │    ││
│  │  └────────────────────────────────────────────────────────────────┘    ││
│  │                               │                                        ││
│  │  ┌────────────────────────────┴───────────────────────────────────┐    ││
│  │  │              EXPLORE/EXPLOIT CONTROLLER                         │    ││
│  │  │   τ_low=0.25 ←──── ADAPTIVE ────→ τ_high=0.60                  │    ││
│  │  └────────────────────────────────────────────────────────────────┘    ││
│  └─────────────────────────────────────────────────────────────────────────┘│
│                                                                              │
│  ┌─────────────────────────────────────────────────────────────────────────┐│
│  │                      KNOWLEDGE GRAPH LAYER                               ││
│  │  ┌───────────────────────────────────────────────────────────────────┐  ││
│  │  │                          Neo4j Database                            │  ││
│  │  │  ┌─────────────────────────────────────────────────────────────┐  │  ││
│  │  │  │ NODES: FINANCIAL_LINE | TEXT_SECTION | NOTE | TABLE_ROW     │  │  ││
│  │  │  └─────────────────────────────────────────────────────────────┘  │  ││
│  │  │  ┌─────────────────────────────────────────────────────────────┐  │  ││
│  │  │  │ EDGES: TEMPORAL_NEXT | REFERS_TO | CAUSED_BY | EXPLAINS     │  │  ││
│  │  │  │        DISCUSSES | SEMANTICALLY_SIMILAR | MENTIONS_ENTITY   │  │  ││
│  │  │  └─────────────────────────────────────────────────────────────┘  │  ││
│  │  └───────────────────────────────────────────────────────────────────┘  ││
│  └─────────────────────────────────────────────────────────────────────────┘│
│                                                                              │
│  ┌─────────────────────────────────────────────────────────────────────────┐│
│  │                      MoE GRAPH BUILDER LAYER                             ││
│  │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐    ││
│  │  │CrossReference│ │CausalChain  │ │Temporal     │ │TableText     │    ││
│  │  │Hunter       │ │Builder      │ │Linker       │ │Connector     │    ││
│  │  └──────────────┘ └──────────────┘ └──────────────┘ └──────────────┘    ││
│  │  ┌──────────────┐ ┌──────────────┐ ┌──────────────────────────────────┐ ││
│  │  │Semantic     │ │Entity       │ │    Connectivity Enforcer          │ ││
│  │  │Bridge       │ │Extractor    │ │    (Bridge Edge Generation)       │ ││
│  │  └──────────────┘ └──────────────┘ └──────────────────────────────────┘ ││
│  └─────────────────────────────────────────────────────────────────────────┘│
│                                                                              │
│  ┌─────────────────────────────────────────────────────────────────────────┐│
│  │                      DATA INGESTION LAYER                                ││
│  │  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐    ││
│  │  │SEC Fetcher  │→│HTML Parser  │→│XBRL         │→│Embedding     │    ││
│  │  │(EDGAR API)  │ │(Sections)   │ │Processor    │ │Engine        │    ││
│  │  └──────────────┘ └──────────────┘ └──────────────┘ └──────────────┘    ││
│  │                                          ↓                               ││
│  │                               ┌──────────────────┐                       ││
│  │                               │FinBERT (768-dim) │                       ││
│  │                               └──────────────────┘                       ││
│  └─────────────────────────────────────────────────────────────────────────┘│
└─────────────────────────────────────────────────────────────────────────────┘
'''

    # Add architecture as monospace text
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)

    doc.add_heading('2.2 Data Flow', level=2)

    doc.add_paragraph(
        'The system processes data through a multi-stage pipeline:'
    )

    flow_steps = [
        ('1. Document Acquisition',
         'SEC EDGAR API fetches Apple 10-K and 10-Q filings (12 documents: 3 annual + 9 quarterly)'),
        ('2. Parsing & Extraction',
         'HTML parser extracts sections (Item 1, 1A, 7, 8), tables, and notes'),
        ('3. XBRL Processing',
         'Financial data extracted with GAAP concept tags and period information'),
        ('4. Chunking',
         'Documents split into manageable chunks (5K-8K characters) with metadata'),
        ('5. Embedding Generation',
         'FinBERT creates 768-dimensional embeddings for each chunk'),
        ('6. Expert Processing',
         'Six experts analyze node pairs to discover typed relationships'),
        ('7. Graph Construction',
         'Nodes and edges loaded into Neo4j with confidence scores'),
        ('8. Connectivity Enforcement',
         'Bridge edges added to ensure single connected component'),
    ]

    for step, desc in flow_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========================================================================
    # SECTION 3: MIXTURE OF EXPERTS FRAMEWORK
    # ========================================================================
    doc.add_heading('3. Mixture of Experts (MoE) Framework', level=1)

    doc.add_paragraph(
        'The MoE framework consists of six specialized experts, each responsible for identifying '
        'specific types of relationships in financial documents. All experts inherit from a common '
        'BaseExpert class that provides standardized interfaces and utility functions.'
    )

    doc.add_heading('3.0 Base Expert Class', level=2)

    doc.add_paragraph(
        'The BaseExpert abstract class (src/experts/base.py) defines the standard interface for all experts:'
    )

    base_methods = [
        ('discover_edges(nodes, embeddings) -> List[Edge]',
         'Abstract method that discovers edges of the expert\'s type. Takes all nodes and their embeddings, returns list of discovered edges with confidence scores.'),
        ('edge_types() -> List[EdgeType]',
         'Abstract method returning the edge types this expert discovers.'),
        ('get_confidence_threshold() -> float',
         'Returns minimum confidence threshold for emitting edges (default: 0.5).'),
        ('evaluate(predictions, gold) -> Metrics',
         'Evaluates expert predictions against gold standard. Returns precision, recall, F1 metrics.'),
        ('_deduplicate_edges(edges) -> List[Edge]',
         'Removes duplicate edges, keeping highest confidence one.'),
        ('_filter_by_confidence(edges, threshold) -> List[Edge]',
         'Filters edges by confidence threshold.'),
        ('_create_edge(...) -> Edge',
         'Helper to create edges with consistent formatting and confidence clamping.'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    hdr_cells[1].text = 'Description'

    for method, desc in base_methods:
        row_cells = table.add_row().cells
        row_cells[0].text = method
        row_cells[1].text = desc

    doc.add_paragraph()

    doc.add_heading('Utility Functions (src/experts/base.py)', level=3)

    util_funcs = [
        ('cosine_similarity(emb1, emb2) -> float',
         'Computes cosine similarity between two embeddings with zero-norm handling.'),
        ('batch_cosine_similarity(query, embeddings) -> np.ndarray',
         'Efficiently computes cosine similarity between a query and multiple embeddings.'),
        ('find_top_k_similar(query_embedding, embeddings, k, threshold, exclude_ids)',
         'Finds top-k most similar nodes to a query, with threshold filtering and exclusion.'),
    ]

    for func, desc in util_funcs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{func}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.1: CrossReferenceHunter
    # ========================================================================
    doc.add_heading('3.1 CrossReferenceHunter Expert', level=2)

    doc.add_paragraph('File: src/experts/cross_reference.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'Detects explicit cross-references in SEC filings such as "See Note 3", '
        '"as discussed in Item 7", or "refer to Part II". This expert identifies '
        'intentional document structure connections that authors explicitly marked.'
    )

    doc.add_heading('Edge Types', level=3)
    edge_types = [
        ('REFERS_TO', 'Explicit reference from source to target (e.g., "See Note 3")'),
        ('EXPLAINS', 'Target explains or provides detail for source'),
    ]

    for etype, desc in edge_types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{etype}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Reference Pattern Categories', level=3)

    patterns_table = doc.add_table(rows=1, cols=3)
    patterns_table.style = 'Table Grid'
    hdr = patterns_table.rows[0].cells
    hdr[0].text = 'Category'
    hdr[1].text = 'Example Patterns'
    hdr[2].text = 'Confidence'

    patterns = [
        ('Note References', 'See Note X, Refer to Note X, Described in Note X', '0.95'),
        ('Item References', 'See Item 7, Discussed in Item 1A', '0.90'),
        ('Part/Item Combined', 'Part II, Item 8', '0.90'),
        ('Section (Quoted)', '"Risk Factors", "Management Discussion"', '0.85'),
        ('Section (Relative)', 'described above, see below', '0.70-0.75'),
        ('Table References', 'See following table, summarized in table', '0.75'),
        ('Financial Statements', 'See Consolidated Statements', '0.80'),
        ('General Notes', 'See accompanying Notes', '0.80'),
    ]

    for cat, pattern, conf in patterns:
        row = patterns_table.add_row().cells
        row[0].text = cat
        row[1].text = pattern
        row[2].text = conf

    doc.add_paragraph()

    doc.add_heading('Algorithm Details', level=3)

    doc.add_paragraph(
        'The CrossReferenceHunter uses a multi-strategy approach:'
    )

    algo_steps = [
        ('Pattern Matching (Primary)',
         'Uses 30+ regex patterns optimized for Apple SEC filing formats. Patterns are ordered '
         'by specificity to avoid duplicate matches at the same position. Each pattern is '
         'associated with a reference type (note, item, section, etc.).'),
        ('LLM Extraction (Optional)',
         'When use_llm=True, the expert uses vLLM to extract additional cross-references that '
         'may be phrased in non-standard ways. LLM results are merged with regex results, '
         'avoiding duplicates.'),
        ('Reference Resolution',
         'Each detected reference is resolved to a target node. Resolution priority: '
         '(1) Same filing exact match, (2) Same filing fuzzy match, (3) Cross-filing match.'),
        ('Confidence Calculation',
         'Base confidence set by reference type. +0.05 bonus for same-filing references.'),
    ]

    for step, desc in algo_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Confidence Score Table', level=3)

    conf_table = doc.add_table(rows=1, cols=2)
    conf_table.style = 'Table Grid'
    hdr = conf_table.rows[0].cells
    hdr[0].text = 'Reference Type'
    hdr[1].text = 'Base Confidence'

    conf_scores = [
        ('note / note_title', '0.95'),
        ('item / part_item', '0.90'),
        ('part', '0.85'),
        ('section_quoted', '0.85'),
        ('notes_general', '0.80'),
        ('financial_statement', '0.80'),
        ('section', '0.75'),
        ('table', '0.75'),
        ('section_relative', '0.70'),
    ]

    for ref_type, conf in conf_scores:
        row = conf_table.add_row().cells
        row[0].text = ref_type
        row[1].text = conf

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.2: CausalChainBuilder
    # ========================================================================
    doc.add_heading('3.2 CausalChainBuilder Expert', level=2)

    doc.add_paragraph('File: src/experts/causal.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'Identifies cause-effect relationships in financial text. Examples include '
        '"Revenue increased due to strong iPhone sales" or "Operating expenses rose, '
        'resulting in lower margins". This expert is critical for answering "why" '
        'questions about financial performance.'
    )

    doc.add_heading('Edge Types', level=3)
    doc.add_paragraph('CAUSED_BY: Effect caused by cause (backward direction)', style='List Bullet')
    doc.add_paragraph('LEADS_TO: Cause leads to effect (forward direction)', style='List Bullet')

    doc.add_heading('Causal Connector Patterns', level=3)

    doc.add_paragraph('Forward Direction Connectors (A leads to B):')
    forward_connectors = [
        'resulted in', 'led to', 'which caused', 'drove [the/higher/lower]',
        'contributed to', 'therefore', 'consequently', 'as a result',
        'leading to', 'resulting in', ', causing'
    ]
    p = doc.add_paragraph()
    p.add_run(', '.join(forward_connectors)).italic = True

    doc.add_paragraph('Backward Direction Connectors (B caused by A):')
    backward_connectors = [
        'due to [the/higher/lower]', 'primarily due to', 'because of',
        'driven by', 'attributed to', 'as a result of', 'resulting from',
        'caused by', 'owing to', 'stemming from', 'offset by',
        'impacted by', 'affected by', 'unfavorable/favorable impact'
    ]
    p = doc.add_paragraph()
    p.add_run(', '.join(backward_connectors)).italic = True

    doc.add_heading('Multi-Strategy Processing', level=3)

    strategies = [
        ('Strategy 1: Direct Node Match',
         'If both cause and effect text match existing nodes, create LEADS_TO edge between them.',
         '0.80'),
        ('Strategy 2: Partial Cause Match',
         'If only cause text matches a node, create CAUSED_BY edge from source to cause node.',
         '0.70'),
        ('Strategy 3: Partial Effect Match',
         'If only effect text matches a node, create LEADS_TO edge from source to effect node.',
         '0.70'),
        ('Strategy 4: Related Financial Node',
         'Find semantically related financial node in same filing using embedding similarity.',
         '0.65'),
        ('Strategy 5: Keyword Fallback',
         'Find financial node with keyword overlap (revenue, margin, expense, etc.).',
         '0.60'),
    ]

    strat_table = doc.add_table(rows=1, cols=3)
    strat_table.style = 'Table Grid'
    hdr = strat_table.rows[0].cells
    hdr[0].text = 'Strategy'
    hdr[1].text = 'Description'
    hdr[2].text = 'Confidence'

    for strat, desc, conf in strategies:
        row = strat_table.add_row().cells
        row[0].text = strat
        row[1].text = desc
        row[2].text = conf

    doc.add_paragraph()

    doc.add_heading('LLM Enhancement', level=3)
    doc.add_paragraph(
        'When use_llm=True, the expert also uses vLLM to extract complex causal relationships '
        'that regex patterns might miss. The LLM is prompted to identify cause, effect, evidence, '
        'confidence, and direction for each relationship. LLM confidence is combined with pattern '
        'confidence for the final score.'
    )

    doc.add_heading('Financial Keywords for Matching', level=3)
    keywords = ['revenue', 'sales', 'profit', 'margin', 'expense', 'cost', 'income',
                'earnings', 'cash', 'debt', 'asset', 'liability', 'iphone', 'mac',
                'ipad', 'services', 'wearables', 'operating', 'gross', 'net', 'interest',
                'tax', 'americas', 'europe', 'china', 'japan', 'asia', 'segment', 'geographic']

    p = doc.add_paragraph()
    p.add_run(', '.join(keywords)).italic = True

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.3: TemporalLinker
    # ========================================================================
    doc.add_heading('3.3 TemporalLinker Expert', level=2)

    doc.add_paragraph('File: src/experts/temporal.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'Links the same financial items across time periods. This enables tracking metrics '
        'like "Revenue FY2022 → Revenue FY2023 → Revenue FY2024" or connecting the same risk '
        'factor discussion across quarterly filings.'
    )

    doc.add_heading('Edge Type', level=3)
    doc.add_paragraph('TEMPORAL_NEXT: Links same entity/metric from earlier to later period', style='List Bullet')

    doc.add_heading('Period Ordering', level=3)
    doc.add_paragraph(
        'The expert maintains a canonical period ordering for chronological linking:'
    )

    periods = [
        'Q1-2022', 'Q2-2022', 'Q3-2022', 'FY2022',
        'Q1-2023', 'Q2-2023', 'Q3-2023', 'FY2023',
        'Q1-2024', 'Q2-2024', 'Q3-2024', 'FY2024'
    ]
    p = doc.add_paragraph()
    p.add_run(' → '.join(periods[:4])).font.size = Pt(9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(' → '.join(periods[4:8])).font.size = Pt(9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(' → '.join(periods[8:])).font.size = Pt(9)

    doc.add_heading('Linking Strategies', level=3)

    temporal_strategies = [
        ('Strategy 1: XBRL Tag Matching',
         'Groups FINANCIAL_LINE nodes by XBRL tag (e.g., us-gaap:Revenues) and creates '
         'edges between consecutive periods. Highest confidence (0.95) because XBRL tags '
         'are standardized identifiers.',
         '0.95'),
        ('Strategy 2: Note Number Matching',
         'Groups NOTE nodes by note number and links same-numbered notes across filings. '
         'High confidence (0.90) because note numbering is consistent.',
         '0.90'),
        ('Strategy 3: Section Matching',
         'Groups TEXT_SECTION nodes by section name (e.g., "Item 7") and links across '
         'filings. Uses embedding similarity (threshold 0.7) to verify content relevance.',
         '0.70+'),
        ('Strategy 4: Embedding Similarity',
         'For TEXT_SECTION and NOTE nodes, finds highest-similarity matches across '
         'consecutive periods. Requires similarity >= 0.90 threshold.',
         '= similarity'),
        ('Strategy 5: LLM Verification',
         'Optional. Pre-filters candidates with embedding similarity >= 0.6, then uses '
         'LLM to verify temporal relationship. Combines LLM confidence with embedding '
         'similarity for final score.',
         'combined'),
    ]

    for strat, desc, conf in temporal_strategies:
        p = doc.add_paragraph()
        p.add_run(f'{strat} (Confidence: {conf})').bold = True
        doc.add_paragraph(desc)

    doc.add_heading('Configuration', level=3)
    config_items = [
        ('similarity_threshold', '0.90', 'Minimum embedding similarity for embedding-based linking'),
        ('use_llm', 'False', 'Whether to use LLM for temporal verification'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for param, default, desc in config_items:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.4: TableTextConnector
    # ========================================================================
    doc.add_heading('3.4 TableTextConnector Expert', level=2)

    doc.add_paragraph('File: src/experts/table_text.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'Connects table rows and financial line items to text that explains them. '
        'For example, linking a revenue table row to the MD&A discussion of revenue '
        'performance, or connecting a balance sheet item to its explanatory note.'
    )

    doc.add_heading('Edge Types', level=3)
    doc.add_paragraph('EXPLAINS_LINE_ITEM: Text explains specific financial line item', style='List Bullet')
    doc.add_paragraph('DISCUSSES: Text discusses the topic of a table/financial item', style='List Bullet')

    doc.add_heading('Numeric Pattern Recognition', level=3)
    doc.add_paragraph(
        'The expert extracts numeric values from text using these patterns:'
    )

    numeric_patterns = [
        ('$X billion/B', 'e.g., "$383 billion" → 383,000,000,000'),
        ('$X million/M', 'e.g., "$200 million" → 200,000,000'),
        ('$X thousand/K', 'e.g., "$50 thousand" → 50,000'),
        ('$X (plain)', 'e.g., "$383,285" → 383,285'),
        ('X percent/%', 'e.g., "25 percent" → 25'),
    ]

    for pattern, example in numeric_patterns:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{pattern}: ').bold = True
        p.add_run(example)

    doc.add_heading('Linking Strategies', level=3)

    table_strategies = [
        ('Strategy 1: Numeric Value Matching',
         'Extracts numeric values from text nodes and matches against FINANCIAL_LINE '
         'or TABLE_ROW node values. Uses 5% tolerance (numeric_tolerance=0.05) to '
         'account for rounding differences.',
         '0.85'),
        ('Strategy 2: XBRL Concept Matching',
         'Parses XBRL tag to extract concept name, converts CamelCase to words, '
         'and searches for keywords in text. Requires ≥50% keyword match AND '
         'embedding similarity ≥0.6.',
         'avg(keyword_ratio, similarity)'),
        ('Strategy 3: Embedding Similarity',
         'Direct embedding similarity between table node and text nodes in same filing. '
         'Uses configurable threshold (default 0.80).',
         '= similarity'),
        ('Strategy 4: LLM Connection',
         'Optional. Groups nodes by filing and uses LLM to identify connections '
         'between table items and narrative text. LLM outputs connection type '
         '(numeric_match, metric_discussion, trend_explanation).',
         'min(llm_conf, sim+0.2)'),
    ]

    for strat, desc, conf in table_strategies:
        p = doc.add_paragraph()
        p.add_run(f'{strat} (Confidence: {conf})').bold = True
        doc.add_paragraph(desc)

    doc.add_heading('Configuration', level=3)
    config_items = [
        ('similarity_threshold', '0.80', 'Minimum embedding similarity threshold'),
        ('numeric_tolerance', '0.05', '5% tolerance for numeric matching'),
        ('use_llm', 'False', 'Enable LLM-based connection finding'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for param, default, desc in config_items:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.5: SemanticBridge
    # ========================================================================
    doc.add_heading('3.5 SemanticBridge Expert', level=2)

    doc.add_paragraph('File: src/experts/semantic.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'Creates edges between semantically similar nodes. Serves as a fallback expert '
        'to ensure graph connectivity when other experts don\'t find explicit relationships. '
        'Uses embedding similarity to identify related content.'
    )

    doc.add_heading('Edge Types', level=3)
    doc.add_paragraph('SEMANTICALLY_SIMILAR: Content is semantically related', style='List Bullet')
    doc.add_paragraph('BRIDGE: Fallback edge for connectivity (lower confidence)', style='List Bullet')

    doc.add_heading('Linking Strategies', level=3)

    semantic_strategies = [
        ('Strategy 1: Within-Filing Similarity',
         'Computes pairwise similarity matrix for all nodes in same filing. '
         'Creates edges for pairs above threshold (0.85). Limits to max 5 edges '
         'per node to avoid over-connectivity. Skips adjacent nodes of same type '
         '(char_offset difference < 1000) as they\'re likely already connected.',
         '= similarity'),
        ('Strategy 2: Cross-Filing Similarity',
         'For each node type, finds best matching node in other filings. '
         'Samples up to 50 nodes per filing for efficiency. Creates edges '
         'for matches above threshold.',
         '= similarity'),
        ('Strategy 3: LLM Semantic Analysis',
         'Optional. Pre-filters candidate pairs with moderate similarity (0.5-0.85) '
         'where regex-based experts might miss relationships. LLM analyzes '
         'relationship type: same_topic, cause_effect, detail_summary, contrast.',
         'avg(llm_sim, emb_sim)'),
    ]

    for strat, desc, conf in semantic_strategies:
        p = doc.add_paragraph()
        p.add_run(f'{strat} (Confidence: {conf})').bold = True
        doc.add_paragraph(desc)

    doc.add_heading('Algorithm: Within-Filing Similarity', level=3)
    doc.add_paragraph(
        'Efficient implementation using matrix operations:'
    )

    algo_steps = [
        'Group nodes by filing_id',
        'Build embedding matrix for each filing',
        'Normalize embeddings (L2 normalization)',
        'Compute similarity matrix: S = E · E^T',
        'For each node i, find j where S[i,j] >= threshold',
        'Limit to max_edges_per_node (default: 5)',
    ]

    for i, step in enumerate(algo_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Configuration', level=3)
    config_items = [
        ('similarity_threshold', '0.85', 'Minimum similarity for SEMANTICALLY_SIMILAR edges'),
        ('bridge_threshold', '0.70', 'Minimum similarity for BRIDGE edges'),
        ('max_edges_per_node', '5', 'Maximum similarity edges per node'),
        ('use_llm', 'False', 'Enable LLM-based semantic analysis'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for param, default, desc in config_items:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    doc.add_page_break()

    # ========================================================================
    # SECTION 3.6: EntityExtractor
    # ========================================================================
    doc.add_heading('3.6 EntityExtractor Expert', level=2)

    doc.add_paragraph('File: src/experts/entity_extractor.py')

    doc.add_heading('Purpose', level=3)
    doc.add_paragraph(
        'LLM-powered expert that extracts named entities from SEC filings and creates '
        'entity nodes with relationships. This enables entity-based queries and connections '
        'like "mentions iPhone" or "discusses Greater China segment".'
    )

    doc.add_heading('Edge Types', level=3)
    doc.add_paragraph('MENTIONS_ENTITY: Document section mentions an entity', style='List Bullet')
    doc.add_paragraph('ENTITY_RELATED_TO: Two entities co-occur in the same context', style='List Bullet')

    doc.add_heading('Entity Types', level=3)

    entity_types = [
        ('COMPANY', 'Companies mentioned (Apple, suppliers, competitors, partners)'),
        ('PRODUCT', 'Products and services (iPhone, Mac, iPad, Apple Watch, Services, AppleCare)'),
        ('SEGMENT', 'Business or geographic segments (Americas, Greater China, Europe)'),
        ('PERSON', 'Key executives or board members mentioned by name'),
        ('FINANCIAL_METRIC', 'Specific financial metrics (Revenue, Net Income, Gross Margin)'),
        ('RISK_FACTOR', 'Identified business risks'),
        ('REGULATION', 'Laws, regulations, accounting standards (GAAP, ASC 606, GDPR)'),
    ]

    et_table = doc.add_table(rows=1, cols=2)
    et_table.style = 'Table Grid'
    hdr = et_table.rows[0].cells
    hdr[0].text = 'Entity Type'
    hdr[1].text = 'Description'

    for etype, desc in entity_types:
        row = et_table.add_row().cells
        row[0].text = etype
        row[1].text = desc

    doc.add_paragraph()

    doc.add_heading('Extraction Process', level=3)

    extraction_steps = [
        ('1. Node Processing',
         'Processes TEXT_SECTION and NOTE nodes (skips FINANCIAL_LINE which are already structured).'),
        ('2. LLM Extraction',
         'For each node, sends text to vLLM with system prompt specifying entity types. '
         'LLM outputs JSON array of {name, type, context} objects.'),
        ('3. Entity Caching',
         'Maintains cache of created entities by (name_lowercase, type) to avoid duplicates. '
         'Same entity mentioned in different sections reuses existing node.'),
        ('4. Edge Creation',
         'Creates MENTIONS_ENTITY edge from source document node to entity node. '
         'Confidence: 0.85'),
        ('5. Co-occurrence Analysis',
         'For entities that appear in the same document section, creates ENTITY_RELATED_TO '
         'edges. Confidence: 0.70'),
    ]

    for step, desc in extraction_steps:
        p = doc.add_paragraph()
        p.add_run(f'{step}: ').bold = True
        p.add_run(desc)

    doc.add_heading('LLM Prompt Template', level=3)

    prompt_text = '''System Prompt:
You are an expert at extracting named entities from Apple SEC filings.
Your task is to identify important entities in financial text.

Entity Types: [COMPANY, PRODUCT, SEGMENT, PERSON, FINANCIAL_METRIC, RISK_FACTOR, REGULATION]

For each entity, output a JSON object with:
- "name": The entity name (use canonical names)
- "type": One of the entity types above
- "context": Brief context about the entity from the text

Output a JSON array. Only extract clearly identifiable entities.'''

    p = doc.add_paragraph()
    p.add_run(prompt_text).font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    doc.add_heading('Requirements', level=3)
    doc.add_paragraph(
        'This expert requires vLLM server to be running (Qwen2.5-7B-Instruct recommended). '
        'Without LLM, the expert returns empty results. Start server with: ./scripts/start_vllm.sh'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 4: ETL PIPELINE
    # ========================================================================
    doc.add_heading('4. ETL Pipeline for Document Ingestion', level=1)

    doc.add_paragraph(
        'The ETL (Extract, Transform, Load) pipeline handles fetching, parsing, and processing '
        'SEC filings into the knowledge graph format. The pipeline is located in src/ingestion/.'
    )

    doc.add_heading('Pipeline Overview', level=2)

    pipeline_diagram = '''
┌──────────────┐    ┌──────────────┐    ┌──────────────┐    ┌──────────────┐
│ SEC Fetcher  │───→│ HTML Parser  │───→│    XBRL      │───→│  Embedding   │
│  (EDGAR API) │    │  (Sections)  │    │  Processor   │    │   Engine     │
└──────────────┘    └──────────────┘    └──────────────┘    └──────────────┘
       │                   │                   │                   │
       ↓                   ↓                   ↓                   ↓
  Raw HTML/XBRL      Parsed Sections      XBRL Values        768-dim Vectors
    Documents          & Tables          & Concepts          (FinBERT)
'''

    p = doc.add_paragraph()
    p.add_run(pipeline_diagram).font.name = 'Courier New'
    p.runs[0].font.size = Pt(9)

    doc.add_heading('4.1 SEC Fetcher Component', level=2)

    doc.add_paragraph('File: src/ingestion/sec_fetcher.py')

    doc.add_paragraph(
        'Fetches Apple 10-K and 10-Q filings from the SEC EDGAR API. Handles rate limiting '
        'and document URL resolution.'
    )

    doc.add_heading('Configuration', level=3)

    sec_config = [
        ('sec_user_agent', 'MoEGraphBuilder research@example.com', 'Required for SEC API access'),
        ('sec_rate_limit', '10.0', 'Requests per second (SEC limit: 10)'),
        ('sec_base_url', 'https://data.sec.gov', 'SEC data API base URL'),
        ('sec_archive_url', 'https://www.sec.gov', 'SEC archive URL for documents'),
        ('apple_cik', '0000320193', 'Apple Inc. CIK number'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for param, default, desc in sec_config:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('Filing Coverage', level=3)

    doc.add_paragraph('Annual Reports (10-K):')
    annual = [
        ('FY2024', '2024-09-28'),
        ('FY2023', '2023-09-30'),
        ('FY2022', '2022-09-24'),
    ]

    for period, date in annual:
        doc.add_paragraph(f'{period} (fiscal year end: {date})', style='List Bullet')

    doc.add_paragraph('Quarterly Reports (10-Q):')
    quarterly = [
        ('Q1-2024: 2023-12-30', 'Q2-2024: 2024-03-30', 'Q3-2024: 2024-06-29'),
        ('Q1-2023: 2022-12-31', 'Q2-2023: 2023-04-01', 'Q3-2023: 2023-07-01'),
        ('Q1-2022: 2021-12-25', 'Q2-2022: 2022-03-26', 'Q3-2022: 2022-06-25'),
    ]

    for row in quarterly:
        doc.add_paragraph(' | '.join(row), style='List Bullet')

    doc.add_page_break()

    doc.add_heading('4.2 HTML Parser Component', level=2)

    doc.add_paragraph('File: src/ingestion/html_parser.py')

    doc.add_paragraph(
        'Parses SEC HTML documents to extract structured sections, tables, and notes.'
    )

    doc.add_heading('Extracted Sections', level=3)

    sections = [
        ('Item 1', 'Business'),
        ('Item 1A', 'Risk Factors'),
        ('Item 7', 'Management Discussion and Analysis (MD&A)'),
        ('Item 8', 'Financial Statements and Supplementary Data'),
        ('Notes', 'Notes to Consolidated Financial Statements'),
    ]

    for item, desc in sections:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{item}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Chunking Configuration', level=3)

    chunk_config = [
        ('chunk_size_default', '5000', 'Default max chars per chunk'),
        ('chunk_size_mda', '8000', 'MD&A (Item 7) - more context needed'),
        ('chunk_size_risk', '8000', 'Risk Factors (Item 1A) - more context needed'),
        ('chunk_min_length', '50', 'Minimum paragraph length to include'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Description'

    for param, val, desc in chunk_config:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = val
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('4.3 XBRL Processor Component', level=2)

    doc.add_paragraph('File: src/ingestion/xbrl_processor.py')

    doc.add_paragraph(
        'Parses XBRL (eXtensible Business Reporting Language) data from SEC filings. '
        'XBRL provides standardized, machine-readable financial data with GAAP concept tags.'
    )

    doc.add_heading('XBRL Concepts Extracted', level=3)

    xbrl_concepts = [
        'GrossProfit', 'CostOfGoodsSold', 'CostOfRevenue',
        'OperatingIncome', 'OperatingExpenses', 'OperatingIncomeLoss',
        'NetIncome', 'NetIncomeLoss', 'ProfitLoss',
        'Revenues', 'RevenueFromContractWithCustomerExcludingAssessedTax',
        'SalesRevenueNet', 'GrossMarginPercentage'
    ]

    p = doc.add_paragraph()
    p.add_run(', '.join(xbrl_concepts)).font.size = Pt(9)

    doc.add_heading('4.4 Embedding Engine (FinBERT)', level=2)

    doc.add_paragraph('File: src/ingestion/embedding_engine.py')

    doc.add_paragraph(
        'Generates domain-specific embeddings using FinBERT (ProsusAI/finbert), a BERT model '
        'pre-trained on financial text. Produces 768-dimensional vectors for semantic similarity.'
    )

    doc.add_heading('Configuration', level=3)

    emb_config = [
        ('finbert_model', 'ProsusAI/finbert', 'Model name/path'),
        ('embedding_batch_size', '32', 'Batch size for processing'),
        ('embedding_max_length', '512', 'Maximum token length'),
        ('device', 'cuda/cpu/mps', 'Compute device'),
    ]

    cfg_table = doc.add_table(rows=1, cols=3)
    cfg_table.style = 'Table Grid'
    hdr = cfg_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Description'

    for param, val, desc in emb_config:
        row = cfg_table.add_row().cells
        row[0].text = param
        row[1].text = val
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('Embedding Process', level=3)

    emb_steps = [
        'Tokenize text with FinBERT tokenizer (truncate to max_length)',
        'Pass through FinBERT model to get hidden states',
        'Apply mean pooling over token dimension',
        'L2 normalize resulting vector',
        'Output: 768-dimensional embedding'
    ]

    for i, step in enumerate(emb_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ========================================================================
    # SECTION 5: COMMUTATOR CONFIGURATION
    # ========================================================================
    doc.add_heading('5. Commutator Configuration', level=1)

    doc.add_paragraph('File: src/opmech/commutator.py')

    doc.add_paragraph(
        'The Commutator is the mathematical core of the OpMech system. It measures the '
        'divergence between two operators to determine when they agree (exploit) or '
        'disagree (explore). This section details the formula and components.'
    )

    doc.add_heading('5.1 Divergence Formula', level=2)

    formula_text = '''
Δ(q, h) = w_E · Δ_E + w_V · Δ_V + w_A · Δ_A + w_C · Δ_C

Where:
    Δ_E = Evidence Divergence (what nodes were retrieved?)
    Δ_V = Structural Divergence (what sections/types were covered?)
    Δ_A = Answer Divergence (how different are the conclusions?)
    Δ_C = Confidence Divergence (how confident is each path?)
'''

    p = doc.add_paragraph()
    p.add_run(formula_text).font.name = 'Courier New'

    doc.add_heading('Default Weights', level=3)

    weights = [
        ('w_E (Evidence)', '0.30', 'Weight for evidence divergence'),
        ('w_V (Structural)', '0.20', 'Weight for structural divergence'),
        ('w_A (Answer)', '0.30', 'Weight for answer divergence'),
        ('w_C (Confidence)', '0.20', 'Weight for confidence divergence'),
    ]

    w_table = doc.add_table(rows=1, cols=3)
    w_table.style = 'Table Grid'
    hdr = w_table.rows[0].cells
    hdr[0].text = 'Weight'
    hdr[1].text = 'Value'
    hdr[2].text = 'Description'

    for weight, val, desc in weights:
        row = w_table.add_row().cells
        row[0].text = weight
        row[1].text = val
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('5.2 Divergence Components', level=2)

    doc.add_heading('Δ_E: Evidence Divergence (Jaccard Distance)', level=3)

    doc.add_paragraph(
        'Measures how different the retrieved evidence sets are between operators.'
    )

    formula = '''
Δ_E = 1 - |E_A ∩ E_B| / |E_A ∪ E_B|

Range: [0, 1]
    - 0 = identical evidence sets
    - 1 = completely disjoint evidence sets
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Δ_V: Structural Divergence', level=3)

    doc.add_paragraph(
        'Measures coverage difference in terms of structural attributes (section, type, period).'
    )

    formula = '''
Δ_V = 1 - |V_A ∩ V_B| / |V_A ∪ V_B|

Where V = {(section, type, period)} tuples

Range: [0, 1]
    - 0 = same structural coverage
    - 1 = completely different structural coverage
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Δ_A: Answer Divergence (Cosine Distance)', level=3)

    doc.add_paragraph(
        'Measures semantic difference between generated answers using FinBERT embeddings.'
    )

    formula = '''
Δ_A = 1 - cos(φ(a_A), φ(a_B))

Where φ is the FinBERT embedding function

Range: [0, 1]
    - 0 = identical semantic meaning
    - 1 = completely opposite meaning
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Δ_C: Confidence Divergence', level=3)

    doc.add_paragraph(
        'Measures disagreement in edge confidence levels between paths.'
    )

    formula = '''
Δ_C = α · |μ_A - μ_B| + β · |σ_A - σ_B| / max(σ_A, σ_B, ε)

Where:
    μ_A, μ_B = mean confidence of edges in path A, B
    σ_A, σ_B = std dev of confidence in path A, B
    α = 0.6 (weight for mean difference)
    β = 0.4 (weight for variance difference)
    ε = 0.01 (avoid division by zero)

Range: [0, 1]
    - 0 = both paths equally confident
    - 1 = one path very confident, other very uncertain
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_page_break()

    doc.add_heading('5.3 Operator Scoring', level=2)

    doc.add_paragraph(
        'The commutator also computes quality scores for each operator\'s output.'
    )

    formula = '''
Score = confidence_score × coverage_score × diversity_score

Where:
    confidence_score = mean(edge_confidences)        [0-1]
    coverage_score = min(1, |evidence| / 10)         [0-1]
    diversity_score = unique_sections / |evidence|   [0-1]

Combined formula:
    Score = confidence × coverage × (0.5 + 0.5 × diversity)
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Score Components Explained', level=3)

    score_components = [
        ('Confidence Score', 'Average edge confidence along the traversal path. High-confidence edges indicate trustworthy relationships.'),
        ('Coverage Score', 'How much evidence was found, normalized to target of 10 nodes. Ensures sufficient grounding.'),
        ('Diversity Score', 'Ratio of unique sections to total evidence. Prevents over-reliance on single section.'),
    ]

    for comp, desc in score_components:
        p = doc.add_paragraph()
        p.add_run(f'{comp}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========================================================================
    # SECTION 6: OPERATOR CONFIGURATIONS
    # ========================================================================
    doc.add_heading('6. Operator Configurations', level=1)

    doc.add_paragraph('File: src/opmech/operators.py')

    doc.add_paragraph(
        'The system uses two parallel operators that traverse the knowledge graph from '
        'different starting points. This dual-operator approach ensures comprehensive '
        'coverage and enables disagreement detection.'
    )

    doc.add_heading('6.1 OperatorA: Structure-First (Numbers → Narrative)', level=2)

    doc.add_heading('Workflow', level=3)

    op_a_steps = [
        'Seed from FINANCIAL_LINE nodes (quantitative data)',
        'Use DIRECT XBRL tag matching for revenue/financial queries',
        'Traverse via structure-oriented edges, WEIGHTED BY EDGE CONFIDENCE',
        'Reach explanatory text through EXPLAINS_LINE_ITEM edges',
    ]

    for i, step in enumerate(op_a_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Query-Specific Features', level=3)

    query_features = [
        ('Revenue Queries',
         'Detects keywords: revenue, sales, net sales, total revenue. '
         'Direct XBRL search for Revenues concept.'),
        ('Expense Queries',
         'Detects keywords: expense, cost, r&d. '
         'Searches for expense/cost XBRL tags.'),
        ('Profit Queries',
         'Detects keywords: profit, income, earnings, margin. '
         'Searches for income XBRL tags.'),
        ('Margin Queries',
         'Special handling using QUERY_TO_XBRL_MAP. '
         'Retrieves GrossProfit, CostOfGoodsSold, etc.'),
        ('Period Filtering',
         'Extracts fiscal year from query (FY2023, 2024, etc.) '
         'and filters results to matching period.'),
    ]

    for feature, desc in query_features:
        p = doc.add_paragraph()
        p.add_run(f'{feature}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Dynamic Edge Selection', level=3)

    doc.add_paragraph(
        'Edge types selected based on explore_weight (w):'
    )

    edge_selection = [
        ('Always Include', 'EXPLAINS_LINE_ITEM, DISCUSSES, TEMPORAL_NEXT, REFERS_TO'),
        ('w > 0.4', '+ CAUSED_BY'),
        ('w > 0.6', '+ MENTIONS_ENTITY'),
        ('Excluded', 'SEMANTICALLY_SIMILAR (penalized in scoring)'),
    ]

    for condition, edges in edge_selection:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{condition}: ').bold = True
        p.add_run(edges)

    doc.add_heading('Evidence Balancing', level=3)

    doc.add_paragraph(
        'Ensures mix of financial and narrative evidence:'
    )

    balance_config = [
        ('min_financial_nodes', '3', 'Ensures XBRL coverage'),
        ('Target Mix', '30% financial + 70% narrative', 'Balanced evidence'),
    ]

    for param, val, desc in balance_config:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{param} = {val}: ').bold = True
        p.add_run(desc)

    doc.add_heading('XBRL Concept Mapping (constants.py)', level=3)

    xbrl_map = '''
QUERY_TO_XBRL_MAP = {
    "gross margin": ["GrossProfit", "CostOfGoodsSold", "CostOfRevenue", "GrossMarginPercentage"],
    "operating margin": ["OperatingIncome", "OperatingExpenses", "OperatingIncomeLoss"],
    "net margin": ["NetIncome", "NetIncomeLoss", "ProfitLoss"],
    "revenue": ["Revenues", "SalesRevenueNet", "RevenueFromContractWithCustomer..."],
    "cost": ["CostOfGoodsSold", "CostOfRevenue", "OperatingExpenses"],
    "earnings": ["NetIncome", "NetIncomeLoss", "ProfitLoss", "EarningsPerShare"],
}
'''
    p = doc.add_paragraph()
    p.add_run(xbrl_map).font.name = 'Courier New'
    p.runs[0].font.size = Pt(8)

    doc.add_page_break()

    doc.add_heading('6.2 OperatorB: Narrative-First (Narrative → Numbers)', level=2)

    doc.add_heading('Workflow', level=3)

    op_b_steps = [
        'Seed from TEXT_SECTION and NOTE nodes (qualitative data)',
        'Traverse via semantic/causal edges, WEIGHTED BY EDGE CONFIDENCE',
        'Reach financial data through DISCUSSES, CAUSED_BY edges',
    ]

    for i, step in enumerate(op_b_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('Dynamic Edge Selection', level=3)

    doc.add_paragraph(
        'Edge types selected based on explore_weight (w):'
    )

    edge_selection = [
        ('Always Include', 'EXPLAINS_LINE_ITEM, DISCUSSES, CAUSED_BY, MENTIONS_ENTITY'),
        ('w < 0.5', '+ TEMPORAL_NEXT (convergence with Operator A)'),
        ('w > 0.6 AND hop == 1', '+ SEMANTICALLY_SIMILAR (limited)'),
        ('Extended', 'REFERS_TO, ENTITY_RELATED_TO, LEADS_TO'),
    ]

    for condition, edges in edge_selection:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{condition}: ').bold = True
        p.add_run(edges)

    doc.add_heading('Evidence Balancing', level=3)

    balance_config = [
        ('min_financial_nodes', '2', 'Less than Operator A'),
        ('Target Mix', '60% narrative + 40% financial', 'Narrative emphasis'),
    ]

    for param, val, desc in balance_config:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{param} = {val}: ').bold = True
        p.add_run(desc)

    doc.add_heading('Evidence Ranking (Both Operators)', level=3)

    doc.add_paragraph(
        'Evidence is ranked by combined score:'
    )

    formula = '''
Combined Score = relevance_weight × similarity + confidence_weight × path_confidence

Where:
    similarity = cosine(query_embedding, node_embedding)
    path_confidence = product of edge confidences along traversal path

This ensures evidence is:
    1. Relevant to the query (high similarity)
    2. Reached via trustworthy paths (high confidence edges)
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Operator Comparison Summary', level=3)

    comp_table = doc.add_table(rows=1, cols=3)
    comp_table.style = 'Table Grid'
    hdr = comp_table.rows[0].cells
    hdr[0].text = 'Aspect'
    hdr[1].text = 'OperatorA (Structure)'
    hdr[2].text = 'OperatorB (Narrative)'

    comparisons = [
        ('Starting Point', 'FINANCIAL_LINE', 'TEXT_SECTION, NOTE'),
        ('Primary Edges', 'TEMPORAL_NEXT, EXPLAINS', 'CAUSED_BY, DISCUSSES'),
        ('XBRL Matching', 'Direct tag matching', 'Via traversal'),
        ('Min Financial', '3 nodes', '2 nodes'),
        ('Evidence Mix', '30% fin / 70% narr', '40% fin / 60% narr'),
        ('Strength', 'Numerical queries', 'Causal/descriptive queries'),
    ]

    for aspect, op_a, op_b in comparisons:
        row = comp_table.add_row().cells
        row[0].text = aspect
        row[1].text = op_a
        row[2].text = op_b

    doc.add_page_break()

    # ========================================================================
    # SECTION 7: EXPLORE/EXPLOIT CONTROLLER
    # ========================================================================
    doc.add_heading('7. Explore/Exploit Controller', level=1)

    doc.add_paragraph('File: src/opmech/controller.py')

    doc.add_paragraph(
        'The controller maps commutator divergence to traversal strategy. It implements '
        'the core explore/exploit tradeoff by adjusting all traversal parameters based '
        'on the divergence level.'
    )

    doc.add_heading('Core Principle', level=2)

    principle = '''
Low divergence (Δ < τ_low = 0.25)  →  EXPLOIT: shallow, focused, confident
High divergence (Δ > τ_high = 0.60) →  EXPLORE: deep, wide, uncertain
Medium divergence                   →  ADAPTIVE: adjust based on trajectory
'''
    p = doc.add_paragraph()
    p.add_run(principle).font.name = 'Courier New'

    doc.add_heading('Explore Weight Calculation', level=2)

    formula = '''
if divergence < τ_low:
    explore_weight = 0.0  # Full exploit
elif divergence > τ_high:
    explore_weight = 1.0  # Full explore
else:
    # Linear interpolation
    explore_weight = (divergence - τ_low) / (τ_high - τ_low)
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('Trajectory Adjustments', level=3)

    adjustments = [
        ('trend > 0.1 (divergence increasing)', 'explore_weight += 0.25 (need more exploration)'),
        ('trend < -0.1 (divergence decreasing)', 'explore_weight -= 0.15 (converging, can exploit)'),
        ('score_diff > 0.3 (one operator much better)', 'explore_weight = max(w, 0.3) (verify)'),
        ('Δ_C > 0.5 (high confidence disagreement)', 'explore_weight += 0.15 (explore more)'),
    ]

    for condition, adjustment in adjustments:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{condition}: ').bold = True
        p.add_run(adjustment)

    doc.add_heading('Parameter Interpolation Table', level=2)

    param_table = doc.add_table(rows=1, cols=4)
    param_table.style = 'Table Grid'
    hdr = param_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Exploit (w=0)'
    hdr[2].text = 'Explore (w=1)'
    hdr[3].text = 'Formula'

    params = [
        ('max_hops', '2', '6', '2 + w × 4'),
        ('seeds_per_operator', '3', '8', '3 + w × 5'),
        ('nodes_per_hop', '5', '15', '5 + w × 10'),
        ('min_edge_confidence', '0.7', '0.4', '0.7 - w × 0.3'),
        ('top_k_evidence', '10', '25', '10 + w × 15'),
        ('confidence_decay', '0.95', '0.85', '0.95 - w × 0.10'),
        ('relevance_weight', '0.5', '0.7', '0.5 + w × 0.2'),
        ('confidence_weight', '0.5', '0.3', '0.5 - w × 0.2'),
    ]

    for param, exploit, explore, formula in params:
        row = param_table.add_row().cells
        row[0].text = param
        row[1].text = exploit
        row[2].text = explore
        row[3].text = formula

    doc.add_paragraph()

    doc.add_heading('Key Insights', level=3)

    insights = [
        ('EXPLOIT mode', 'High min_edge_confidence, balance relevance/confidence equally, fewer hops'),
        ('EXPLORE mode', 'Low min_edge_confidence (accept weaker edges), prioritize relevance, more hops'),
        ('Confidence Decay', 'Path confidence = Π(edge_conf × decay^hop). Exploit uses higher decay (0.95) to trust longer paths.'),
    ]

    for mode, insight in insights:
        p = doc.add_paragraph()
        p.add_run(f'{mode}: ').bold = True
        p.add_run(insight)

    doc.add_heading('Edge Type Selection by Mode', level=3)

    edge_table = doc.add_table(rows=1, cols=3)
    edge_table.style = 'Table Grid'
    hdr = edge_table.rows[0].cells
    hdr[0].text = 'w Range'
    hdr[1].text = 'Operator A Edges'
    hdr[2].text = 'Operator B Edges'

    edge_modes = [
        ('w < 0.3', 'TEMPORAL_NEXT, EXPLAINS_LINE_ITEM', 'CAUSED_BY, DISCUSSES'),
        ('0.3 ≤ w < 0.7', '+ REFERS_TO', '+ MENTIONS_ENTITY, REFERS_TO'),
        ('w ≥ 0.7', '+ DISCUSSES, SEMANTICALLY_SIMILAR', '+ SEM_SIMILAR, ENTITY_RELATED, LEADS_TO'),
    ]

    for w_range, op_a, op_b in edge_modes:
        row = edge_table.add_row().cells
        row[0].text = w_range
        row[1].text = op_a
        row[2].text = op_b

    doc.add_page_break()

    # ========================================================================
    # SECTION 8: MODE SELECTION
    # ========================================================================
    doc.add_heading('8. Mode Selection & Trust Decisions', level=1)

    doc.add_paragraph('File: src/opmech/mode_selection.py')

    doc.add_paragraph(
        'The Mode Selector determines the final output mode (EXPLOIT/ADAPTIVE/EXPLORE) '
        'and decides which operator to trust when they disagree. This is critical for '
        'producing reliable answers.'
    )

    doc.add_heading('8.1 Query Modes', level=2)

    modes = [
        ('EXPLOIT', 'High confidence, clear answer. Both operators agree or one is clearly more reliable.'),
        ('ADAPTIVE', 'Moderate confidence, balanced view. Merge answers with appropriate weighting.'),
        ('EXPLORE', 'Low confidence, multiple perspectives. Present both viewpoints with uncertainty.'),
    ]

    for mode, desc in modes:
        p = doc.add_paragraph()
        p.add_run(f'{mode}: ').bold = True
        p.add_run(desc)

    doc.add_heading('8.2 Trust Decisions', level=2)

    trust_decisions = [
        ('TRUST_A', 'Structure-first operator is more reliable (typically for numerical queries)'),
        ('TRUST_B', 'Narrative-first operator is more reliable (typically for causal queries)'),
        ('MERGE_EQUAL', 'Both equally reliable, merge with equal weight'),
        ('MERGE_WEIGHTED', 'Merge with reliability-based weighting'),
        ('CONFLICT', 'Irreconcilable conflict, present both perspectives'),
    ]

    for decision, desc in trust_decisions:
        p = doc.add_paragraph()
        p.add_run(f'{decision}: ').bold = True
        p.add_run(desc)

    doc.add_heading('8.3 Operator Reliability Scoring', level=2)

    doc.add_heading('Source Authority Scores', level=3)

    authority_table = doc.add_table(rows=1, cols=2)
    authority_table.style = 'Table Grid'
    hdr = authority_table.rows[0].cells
    hdr[0].text = 'Node Type'
    hdr[1].text = 'Authority Score'

    authorities = [
        ('FINANCIAL_LINE', '1.0 (XBRL-tagged, audited)'),
        ('TABLE_ROW', '0.9 (Structured data)'),
        ('TABLE', '0.85'),
        ('TEXT_SECTION', '0.6 (Narrative, contextual)'),
        ('NOTE', '0.5 (May reference multiple periods)'),
        ('ENTITY', '0.4 (Needs context)'),
    ]

    for node_type, score in authorities:
        row = authority_table.add_row().cells
        row[0].text = node_type
        row[1].text = score

    doc.add_paragraph()

    doc.add_heading('Query-Evidence Fit Scores', level=3)

    doc.add_paragraph(
        'Different query types need different evidence types:'
    )

    fit_table = doc.add_table(rows=1, cols=3)
    fit_table.style = 'Table Grid'
    hdr = fit_table.rows[0].cells
    hdr[0].text = 'Query Type'
    hdr[1].text = 'Best Evidence'
    hdr[2].text = 'Fit Score'

    fits = [
        ('NUMERICAL', 'FINANCIAL_LINE', '1.0'),
        ('NUMERICAL', 'TABLE_ROW', '0.9'),
        ('NUMERICAL', 'TEXT_SECTION', '0.4'),
        ('CAUSAL', 'TEXT_SECTION', '1.0'),
        ('CAUSAL', 'NOTE', '0.8'),
        ('CAUSAL', 'FINANCIAL_LINE', '0.4'),
        ('DESCRIPTIVE', 'TEXT_SECTION', '1.0'),
        ('DESCRIPTIVE', 'NOTE', '0.9'),
    ]

    for query_type, evidence, fit in fits:
        row = fit_table.add_row().cells
        row[0].text = query_type
        row[1].text = evidence
        row[2].text = fit

    doc.add_paragraph()

    doc.add_heading('Reliability Calculation', level=3)

    formula = '''
reliability = w_eq × evidence_quality + w_sa × source_authority +
              w_qf × query_fit + w_pc × path_confidence

Weights by query type:
    NUMERICAL: w_eq=0.15, w_sa=0.40, w_qf=0.30, w_pc=0.15
    CAUSAL:    w_eq=0.20, w_sa=0.15, w_qf=0.45, w_pc=0.20
    BALANCED:  w_eq=0.25, w_sa=0.25, w_qf=0.30, w_pc=0.20
'''
    p = doc.add_paragraph()
    p.add_run(formula).font.name = 'Courier New'

    doc.add_heading('8.4 EXPLOIT Conditions', level=2)

    exploit_conditions = [
        'Strong answer agreement (Δ_A < 0.15)',
        'Clear reliability winner with trusted_reliability > 0.70',
        'Good overall convergence (Δ < 0.35) with converging/stable trajectory',
        'Simple query with moderate agreement',
    ]

    doc.add_paragraph('EXPLOIT requires ≥2 conditions:')
    for cond in exploit_conditions:
        doc.add_paragraph(cond, style='List Bullet')

    doc.add_heading('8.5 EXPLORE Conditions (Hard Triggers)', level=2)

    explore_conditions = [
        'Strong answer disagreement (Δ_A > 0.40) with no clear winner',
        'Diverging trajectory over multiple hops',
        'Opinion/speculative query type',
        'Both operators unreliable (reliability < 0.45)',
        'Very high divergence (Δ > 0.60 AND Δ_E > 0.80)',
    ]

    doc.add_paragraph('EXPLORE requires only 1 condition:')
    for cond in explore_conditions:
        doc.add_paragraph(cond, style='List Bullet')

    doc.add_heading('8.6 Financial Dominance Threshold', level=2)

    doc.add_paragraph(
        'Critical fix for numerical queries: If one operator has ≥55% of FINANCIAL_LINE '
        'evidence, trust that operator. This prevents merging to incorrect answers when '
        'XBRL data clearly supports one conclusion.'
    )

    doc.add_page_break()

    # ========================================================================
    # SECTION 9: SYSTEM INTEGRATION
    # ========================================================================
    doc.add_heading('9. System Integration & Query Flow', level=1)

    doc.add_paragraph('File: src/opmech/system.py')

    doc.add_paragraph(
        'The OpMechGraphRAG class is the main orchestrator that integrates all components '
        'into a cohesive query processing system.'
    )

    doc.add_heading('9.1 Initialization', level=2)

    init_params = [
        ('neo4j_uri', 'bolt://localhost:7687', 'Neo4j connection URI'),
        ('neo4j_user', 'neo4j', 'Database username'),
        ('neo4j_password', 'password123', 'Database password'),
        ('vllm_url', 'http://localhost:8000/v1', 'vLLM server URL'),
        ('tau_low', '0.25', 'Low divergence threshold'),
        ('tau_high', '0.60', 'High divergence threshold'),
        ('max_hops', '6', 'Maximum traversal hops'),
    ]

    init_table = doc.add_table(rows=1, cols=3)
    init_table.style = 'Table Grid'
    hdr = init_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for param, default, desc in init_params:
        row = init_table.add_row().cells
        row[0].text = param
        row[1].text = default
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('Initialized Components', level=3)

    components = [
        'KnowledgeGraphInterface (Neo4j client)',
        'LLMInterface (vLLM client)',
        'FinBERT embedding function',
        'OperatorA and OperatorB',
        'ExploreExploitController',
        'ModeSelector',
        'HybridQueryClassifier',
        'RobustConsistencyChecker',
    ]

    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    doc.add_heading('9.2 Query Processing Flow', level=2)

    doc.add_heading('Phase 1: Query Classification', level=3)

    flow_text = '''
Query → Classify (HybridQueryClassifier)
    ├─ Query Type: NUMERICAL | TEMPORAL | CAUSAL | DESCRIPTIVE | OPINION
    ├─ Complexity: simple | moderate | complex
    ├─ Numerical Expected: bool
    └─ Effective Max Hops: 3-7 based on complexity
'''
    p = doc.add_paragraph()
    p.add_run(flow_text).font.name = 'Courier New'

    doc.add_heading('Phase 2: Hop Iteration', level=3)

    hop_flow = '''
for hop in 1..max_hops:

    # Two-phase operator execution
    if hop == 1:
        # Independent exploration
        belief_A = OperatorA.execute(query, strategy, other_evidence=None)
        belief_B = OperatorB.execute(query, strategy, other_evidence=None)
    else:
        # Convergence-aware re-exploration
        belief_A = OperatorA.execute(query, strategy, other_evidence=B_evidence)
        belief_B = OperatorB.execute(query, strategy, other_evidence=A_evidence)

    # Generate answers via LLM
    belief_A.answer = LLM.generate_answer(query, belief_A.evidence)
    belief_B.answer = LLM.generate_answer(query, belief_B.evidence)

    # Compute commutator
    commutator = compute_commutator(belief_A, belief_B, embed_fn)
    trajectory.append(commutator)

    # Apply convergence pressure if Δ_E > 0.8
    apply_convergence_pressure(belief_A, belief_B, commutator.delta_E)

    # Update strategy
    strategy = controller.compute_strategy(commutator, trajectory)

    # Check termination
    if should_terminate(trajectory, hop, max_hops, query_class):
        break
'''
    p = doc.add_paragraph()
    p.add_run(hop_flow).font.name = 'Courier New'
    p.runs[0].font.size = Pt(8)

    doc.add_heading('Phase 3: Final Answer Generation', level=3)

    final_flow = '''
# Determine mode with full context
mode_decision = mode_selector.determine_mode(
    commutator, trajectory, query,
    evidence_types_A, evidence_types_B,
    path_confidence_A, path_confidence_B
)

# Check consistency between operators
consistency = consistency_checker.check_consistency(output_A, output_B)

# Generate answer based on mode and trust
answer = generate_answer_with_trust(
    query, belief_A, belief_B, mode_decision, trajectory
)

# Return QueryResult with full diagnostics
'''
    p = doc.add_paragraph()
    p.add_run(final_flow).font.name = 'Courier New'

    doc.add_heading('9.3 Termination Conditions', level=2)

    termination = [
        ('Reached max hops', 'Safety limit reached'),
        ('Strong convergence', 'Δ < τ_low (0.25)'),
        ('Excellent answer agreement', 'Δ_A < 0.05 for numerical queries with Δ < 0.40'),
        ('Opinion query minimum', 'At least 3 hops for opinion queries'),
        ('Stabilized', 'Improvement < 0.02 for 2+ hops with Δ < 0.45'),
        ('Simple numerical converged', 'Δ < 0.35 AND Δ_A < 0.10'),
        ('Diverging rapidly', '3 consecutive hops with increasing divergence'),
    ]

    for condition, desc in termination:
        p = doc.add_paragraph()
        p.add_run(f'{condition}: ').bold = True
        p.add_run(desc)

    doc.add_heading('9.4 Convergence Pressure Mechanism', level=2)

    doc.add_paragraph(
        'When operators diverge too much (Δ_E > 0.8), the system shares top evidence nodes '
        'between operators as "bridge seeds" to encourage convergence.'
    )

    pressure_code = '''
if delta_E > CONVERGENCE_PRESSURE_THRESHOLD (0.8):
    # Top nodes from A → give to B as bridge seeds
    bridge_A_to_B = belief_A.evidence[:3].ids
    operator_B.add_bridge_seeds(bridge_A_to_B)

    # Top nodes from B → give to A as bridge seeds
    bridge_B_to_A = belief_B.evidence[:3].ids
    operator_A.add_bridge_seeds(bridge_B_to_A)
'''
    p = doc.add_paragraph()
    p.add_run(pressure_code).font.name = 'Courier New'

    doc.add_page_break()

    # ========================================================================
    # SECTION 10: DATA CLASSES
    # ========================================================================
    doc.add_heading('10. Data Classes & Models', level=1)

    doc.add_paragraph('Files: src/opmech/data_classes.py, src/models.py')

    doc.add_heading('10.1 Core Data Classes (opmech/data_classes.py)', level=2)

    doc.add_heading('Node', level=3)
    node_fields = [
        ('id', 'str', 'Unique identifier'),
        ('type', 'str', 'Node type (FINANCIAL_LINE, TEXT_SECTION, etc.)'),
        ('text', 'str', 'Text content'),
        ('metadata', 'Dict[str, Any]', 'Additional metadata'),
        ('embedding', 'np.ndarray', '768-dim embedding (optional)'),
    ]

    node_table = doc.add_table(rows=1, cols=3)
    node_table.style = 'Table Grid'
    hdr = node_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in node_fields:
        row = node_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('Edge', level=3)
    edge_fields = [
        ('source_id', 'str', 'Source node ID'),
        ('target_id', 'str', 'Target node ID'),
        ('type', 'str', 'Edge type (REFERS_TO, CAUSED_BY, etc.)'),
        ('confidence', 'float', 'Confidence score [0.0-1.0]'),
        ('expert', 'str', 'Expert that created this edge'),
        ('evidence', 'str', 'Text evidence for relationship'),
    ]

    edge_table = doc.add_table(rows=1, cols=3)
    edge_table.style = 'Table Grid'
    hdr = edge_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in edge_fields:
        row = edge_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('BeliefState', level=3)
    belief_fields = [
        ('evidence', 'List[Node]', 'Retrieved evidence nodes'),
        ('answer', 'str', 'Generated answer text'),
        ('edge_confidences', 'List[float]', 'All edge confidences from traversal'),
        ('evidence_confidences', 'List[float]', 'Per-evidence path confidence'),
        ('mean_path_confidence', 'float', 'Aggregate path confidence'),
        ('operator_path', 'str', '"structure_first" or "narrative_first"'),
        ('hops_used', 'int', 'Number of hops in traversal'),
        ('seeds_used', 'List[str]', 'Seed node IDs'),
        ('edges_traversed', 'List[Edge]', 'Edges traversed'),
    ]

    belief_table = doc.add_table(rows=1, cols=3)
    belief_table.style = 'Table Grid'
    hdr = belief_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in belief_fields:
        row = belief_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('TraversalStrategy', level=3)
    strat_fields = [
        ('max_hops', 'int', 'Maximum traversal depth (1-6)'),
        ('current_hop', 'int', 'Current hop in iteration'),
        ('seeds_per_operator', 'int', 'Initial seed nodes (3-10)'),
        ('nodes_per_hop', 'int', 'Max nodes per hop (5-20)'),
        ('edge_types_A', 'List[str]', 'Edge types for Operator A'),
        ('edge_types_B', 'List[str]', 'Edge types for Operator B'),
        ('min_edge_confidence', 'float', 'Minimum edge confidence (0.3-0.8)'),
        ('top_k_evidence', 'int', 'Final evidence selection (10-30)'),
        ('confidence_decay', 'float', 'Decay factor per hop (0.85-0.95)'),
        ('relevance_weight', 'float', 'Weight for similarity (0.4-0.7)'),
        ('confidence_weight', 'float', 'Weight for path confidence (0.3-0.6)'),
        ('output_mode', 'str', '"exploit", "adaptive", "explore"'),
        ('explore_weight', 'float', '0.0=exploit, 1.0=explore'),
    ]

    strat_table = doc.add_table(rows=1, cols=3)
    strat_table.style = 'Table Grid'
    hdr = strat_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in strat_fields:
        row = strat_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_page_break()

    doc.add_heading('CommutatorResult', level=3)
    comm_fields = [
        ('delta_E', 'float', 'Evidence divergence'),
        ('delta_V', 'float', 'Structural divergence'),
        ('delta_A', 'float', 'Answer divergence'),
        ('delta_C', 'float', 'Confidence divergence'),
        ('combined', 'float', 'Weighted combined divergence'),
        ('weights', 'Dict[str, float]', 'Component weights'),
        ('hop', 'int', 'Hop number'),
        ('operator_A_score', 'float', 'Operator A quality score'),
        ('operator_B_score', 'float', 'Operator B quality score'),
    ]

    comm_table = doc.add_table(rows=1, cols=3)
    comm_table.style = 'Table Grid'
    hdr = comm_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in comm_fields:
        row = comm_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('QueryResult', level=3)
    result_fields = [
        ('answer', 'str', 'Final answer text'),
        ('confidence', 'float', 'Overall confidence'),
        ('mode', 'OutputMode', 'EXPLOIT/ADAPTIVE/EXPLORE'),
        ('hops_used', 'int', 'Total hops executed'),
        ('trajectory', 'List[CommutatorResult]', 'Divergence history'),
        ('evidence_A', 'List[Node]', 'Operator A evidence'),
        ('evidence_B', 'List[Node]', 'Operator B evidence'),
        ('answer_A', 'str', 'Operator A answer'),
        ('answer_B', 'str', 'Operator B answer'),
        ('reasoning', 'str', 'Mode decision reasoning'),
        ('operator_scores', 'Dict[str, float]', 'Operator quality scores'),
        ('path_confidence_A', 'float', 'Operator A path confidence'),
        ('path_confidence_B', 'float', 'Operator B path confidence'),
        ('edge_conf_stats', 'Dict', 'Edge confidence statistics'),
    ]

    result_table = doc.add_table(rows=1, cols=3)
    result_table.style = 'Table Grid'
    hdr = result_table.rows[0].cells
    hdr[0].text = 'Field'
    hdr[1].text = 'Type'
    hdr[2].text = 'Description'

    for field, ftype, desc in result_fields:
        row = result_table.add_row().cells
        row[0].text = field
        row[1].text = ftype
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('10.2 Enumerations', level=2)

    doc.add_heading('NodeType', level=3)
    node_types = ['FINANCIAL_LINE', 'TEXT_SECTION', 'NOTE', 'TABLE_ROW', 'ENTITY']
    doc.add_paragraph(', '.join(node_types))

    doc.add_heading('EdgeType', level=3)
    edge_types = [
        'EXPLAINS', 'REFERS_TO', 'CAUSED_BY', 'LEADS_TO', 'TEMPORAL_NEXT',
        'EXPLAINS_LINE_ITEM', 'DISCUSSES', 'SEMANTICALLY_SIMILAR', 'BRIDGE',
        'MENTIONS_ENTITY', 'ENTITY_RELATED_TO'
    ]
    doc.add_paragraph(', '.join(edge_types))

    doc.add_heading('OutputMode', level=3)
    doc.add_paragraph('EXPLOIT | ADAPTIVE | EXPLORE')

    doc.add_heading('QueryMode', level=3)
    doc.add_paragraph('EXPLOIT | ADAPTIVE | EXPLORE')

    doc.add_heading('TrustDecision', level=3)
    doc.add_paragraph('TRUST_A | TRUST_B | MERGE_EQUAL | MERGE_WEIGHTED | CONFLICT')

    doc.add_heading('QueryType', level=3)
    doc.add_paragraph('NUMERICAL | TEMPORAL | CAUSAL | DESCRIPTIVE | OPINION')

    doc.add_page_break()

    # ========================================================================
    # SECTION 11: CONFIGURATION
    # ========================================================================
    doc.add_heading('11. Configuration Reference', level=1)

    doc.add_paragraph('File: src/config.py')

    doc.add_heading('Environment Variables (.env)', level=2)

    env_vars = [
        ('NEO4J_URI', 'bolt://localhost:7687', 'Neo4j connection URI'),
        ('NEO4J_USER', 'neo4j', 'Neo4j username'),
        ('NEO4J_PASSWORD', 'password123', 'Neo4j password'),
        ('VLLM_API_BASE', 'http://localhost:8000/v1', 'vLLM API URL'),
        ('VLLM_MODEL', 'Qwen/Qwen2.5-7B-Instruct', 'LLM model'),
        ('DEVICE', 'cuda', 'Compute device (cuda/cpu/mps)'),
    ]

    env_table = doc.add_table(rows=1, cols=3)
    env_table.style = 'Table Grid'
    hdr = env_table.rows[0].cells
    hdr[0].text = 'Variable'
    hdr[1].text = 'Default'
    hdr[2].text = 'Description'

    for var, default, desc in env_vars:
        row = env_table.add_row().cells
        row[0].text = var
        row[1].text = default
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('Expert Thresholds', level=2)

    thresholds = [
        ('cross_ref_confidence_threshold', '0.5', 'CrossReferenceHunter'),
        ('causal_confidence_threshold', '0.5', 'CausalChainBuilder'),
        ('temporal_similarity_threshold', '0.90', 'TemporalLinker'),
        ('table_text_similarity_threshold', '0.80', 'TableTextConnector'),
        ('semantic_similarity_threshold', '0.85', 'SemanticBridge'),
        ('bridge_similarity_threshold', '0.70', 'Connectivity Enforcer'),
    ]

    thresh_table = doc.add_table(rows=1, cols=3)
    thresh_table.style = 'Table Grid'
    hdr = thresh_table.rows[0].cells
    hdr[0].text = 'Setting'
    hdr[1].text = 'Value'
    hdr[2].text = 'Expert'

    for setting, value, expert in thresholds:
        row = thresh_table.add_row().cells
        row[0].text = setting
        row[1].text = value
        row[2].text = expert

    doc.add_paragraph()

    doc.add_heading('System Thresholds', level=2)

    sys_thresholds = [
        ('tau_low', '0.25', 'Exploit below this divergence'),
        ('tau_high', '0.60', 'Explore above this divergence'),
        ('max_hops', '6', 'Maximum traversal hops'),
        ('CONVERGENCE_PRESSURE_THRESHOLD', '0.80', 'Share nodes above this Δ_E'),
        ('FINANCIAL_DOMINANCE_THRESHOLD', '0.55', 'Trust operator with this % XBRL'),
        ('min_improvement', '0.02', 'Minimum delta improvement to continue'),
        ('min_hops_opinion', '3', 'Minimum hops for opinion queries'),
    ]

    sys_table = doc.add_table(rows=1, cols=3)
    sys_table.style = 'Table Grid'
    hdr = sys_table.rows[0].cells
    hdr[0].text = 'Parameter'
    hdr[1].text = 'Value'
    hdr[2].text = 'Description'

    for param, value, desc in sys_thresholds:
        row = sys_table.add_row().cells
        row[0].text = param
        row[1].text = value
        row[2].text = desc

    doc.add_page_break()

    # ========================================================================
    # SECTION 12: APPENDICES
    # ========================================================================
    doc.add_heading('12. Appendices', level=1)

    doc.add_heading('Appendix A: File Structure', level=2)

    file_structure = '''
/src/
├── experts/                    # MoE Experts
│   ├── __init__.py            # Expert factory functions
│   ├── base.py                # Abstract base class
│   ├── semantic.py            # SemanticBridge expert
│   ├── temporal.py            # TemporalLinker expert
│   ├── causal.py              # CausalChainBuilder expert
│   ├── cross_reference.py     # CrossReferenceHunter expert
│   ├── table_text.py          # TableTextConnector expert
│   ├── entity_extractor.py    # EntityExtractor expert
│   └── llm_client.py          # vLLM client for LLM features
│
├── ingestion/                  # ETL Pipeline
│   ├── sec_fetcher.py         # SEC EDGAR API client
│   ├── html_parser.py         # HTML document parser
│   ├── xbrl_processor.py      # XBRL data extractor
│   └── embedding_engine.py    # FinBERT embedding generator
│
├── graph/                      # Graph Construction
│   ├── builder.py             # Main graph builder orchestrator
│   ├── connectivity.py        # Bridge edge generation
│   └── neo4j_client.py        # Neo4j database interface
│
├── opmech/                     # Operator Mechanics
│   ├── system.py              # Main OpMechGraphRAG class
│   ├── operators.py           # OperatorA and OperatorB
│   ├── commutator.py          # Divergence computation
│   ├── controller.py          # Explore/exploit controller
│   ├── mode_selection.py      # Mode and trust decisions
│   ├── query_classifier.py    # Query type classification
│   ├── data_classes.py        # Core data structures
│   ├── constants.py           # XBRL mappings
│   ├── graph_interface.py     # Graph query interface
│   ├── llm_interface.py       # Answer generation
│   └── robust_consistency_checker.py  # Consistency validation
│
├── config.py                   # Application settings
└── models.py                   # Pydantic data models
'''

    p = doc.add_paragraph()
    p.add_run(file_structure).font.name = 'Courier New'
    p.runs[0].font.size = Pt(8)

    doc.add_heading('Appendix B: Edge Type Reference', level=2)

    edge_ref = [
        ('TEMPORAL_NEXT', 'Links same metric across time periods', 'TemporalLinker', '0.70-0.95'),
        ('REFERS_TO', 'Explicit cross-reference', 'CrossReferenceHunter', '0.70-0.95'),
        ('EXPLAINS', 'Target explains source', 'CrossReferenceHunter', '0.70-0.95'),
        ('CAUSED_BY', 'Effect caused by cause', 'CausalChainBuilder', '0.60-0.80'),
        ('LEADS_TO', 'Cause leads to effect', 'CausalChainBuilder', '0.60-0.80'),
        ('EXPLAINS_LINE_ITEM', 'Text explains financial item', 'TableTextConnector', '0.65-0.85'),
        ('DISCUSSES', 'Text discusses topic', 'TableTextConnector', '0.65-0.85'),
        ('SEMANTICALLY_SIMILAR', 'High embedding similarity', 'SemanticBridge', '= similarity'),
        ('BRIDGE', 'Connectivity enforcement', 'ConnectivityEnforcer', '0.50-0.70'),
        ('MENTIONS_ENTITY', 'Document mentions entity', 'EntityExtractor', '0.85'),
        ('ENTITY_RELATED_TO', 'Entities co-occur', 'EntityExtractor', '0.70'),
    ]

    edge_ref_table = doc.add_table(rows=1, cols=4)
    edge_ref_table.style = 'Table Grid'
    hdr = edge_ref_table.rows[0].cells
    hdr[0].text = 'Edge Type'
    hdr[1].text = 'Description'
    hdr[2].text = 'Created By'
    hdr[3].text = 'Confidence Range'

    for etype, desc, expert, conf in edge_ref:
        row = edge_ref_table.add_row().cells
        row[0].text = etype
        row[1].text = desc
        row[2].text = expert
        row[3].text = conf

    doc.add_paragraph()

    doc.add_heading('Appendix C: Startup Commands', level=2)

    commands = '''
# Start Neo4j
docker run -d --name neo4j \\
  -p 7474:7474 -p 7687:7687 \\
  -e NEO4J_AUTH=neo4j/password123 \\
  neo4j:latest

# Start vLLM server
./scripts/start_vllm.sh Qwen/Qwen2.5-7B-Instruct 8000

# Build knowledge graph
python -m src.graph.builder

# Run queries
python -m src.opmech.system --query "What was Apple's revenue in FY2024?"
'''

    p = doc.add_paragraph()
    p.add_run(commands).font.name = 'Courier New'

    # Save document
    doc.save('/home/divyansh/AIF_FInal_Project/docs/MOE_Architecture_Technical_Documentation.docx')
    print("Document saved: docs/MOE_Architecture_Technical_Documentation.docx")

if __name__ == '__main__':
    create_document()
